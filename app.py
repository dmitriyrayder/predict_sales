
import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from prophet import Prophet
from sklearn.metrics import mean_absolute_error, mean_squared_error, mean_absolute_percentage_error, r2_score
from scipy.signal import savgol_filter
from io import BytesIO
from datetime import datetime
import warnings
warnings.filterwarnings('ignore')

# Конфигурация страницы
st.set_page_config(
    page_title="🏪 Система прогнозирования продаж",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS стили
st.markdown("""
<style>
    .main-header {
        font-size: 3rem;
        font-weight: bold;
        background: linear-gradient(120deg, #1f77b4, #667eea);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        text-align: center;
        margin-bottom: 2rem;
        text-shadow: 2px 2px 4px rgba(0,0,0,0.1);
    }
    
    .metric-container {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 1.5rem;
        border-radius: 15px;
        color: white;
        text-align: center;
        margin: 0.5rem 0;
        box-shadow: 0 8px 16px rgba(0,0,0,0.2);
        transition: transform 0.3s ease;
    }
    
    .metric-container:hover {
        transform: translateY(-5px);
        box-shadow: 0 12px 24px rgba(0,0,0,0.3);
    }
    
    .insight-card {
        background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%);
        padding: 1.2rem;
        border-radius: 12px;
        color: white;
        margin: 0.5rem 0;
        box-shadow: 0 6px 12px rgba(0,0,0,0.15);
        border-left: 5px solid #fff;
    }
    
    .problem-card {
        background: linear-gradient(135deg, #ff6b6b 0%, #feca57 100%);
        padding: 1.2rem;
        border-radius: 12px;
        margin: 0.5rem 0;
        border-left: 5px solid #ee5a6f;
        box-shadow: 0 6px 12px rgba(0,0,0,0.15);
        color: white;
        font-weight: 500;
    }
    
    .accuracy-card {
        background: linear-gradient(135deg, #43e97b 0%, #38f9d7 100%);
        padding: 1.5rem;
        border-radius: 15px;
        margin: 1rem 0;
        box-shadow: 0 8px 16px rgba(0,0,0,0.2);
        color: white;
    }
</style>
""", unsafe_allow_html=True)

@st.cache_data
def load_and_validate_data(uploaded_file):
    """Загружает и валидирует данные из Excel файла"""
    try:
        progress_bar = st.progress(0)
        progress_bar.progress(25)
        
        df = pd.read_excel(uploaded_file)
        progress_bar.progress(50)
        
        required_cols = ['Magazin', 'Datasales', 'Art', 'Describe', 'Model', 'Segment', 'Price', 'Qty', 'Sum']
        missing_cols = [col for col in required_cols if col not in df.columns]
        
        if missing_cols:
            st.error(f"❌ Отсутствуют обязательные колонки: {missing_cols}")
            return None
            
        progress_bar.progress(75)
        
        df['Datasales'] = pd.to_datetime(df['Datasales'], errors='coerce', dayfirst=True)
        df = df.dropna(subset=['Datasales']).sort_values('Datasales')
        df = df[(df['Qty'] >= 0) & (df['Price'] > 0)]
        
        progress_bar.progress(100)
        progress_bar.empty()
        
        st.success(f"✅ Данные успешно загружены! Обработано {len(df)} записей")
        return df
        
    except Exception as e:
        st.error(f"❌ Ошибка при загрузке файла: {str(e)}")
        return None

def show_data_statistics(df):
    """Отображает статистику данных"""
    st.markdown("## 📊 Статистика данных")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.markdown(
            f"""<div class="metric-container">
                <h3>📦 Всего записей</h3>
                <h2>{len(df):,}</h2>
            </div>""", 
            unsafe_allow_html=True
        )
        
    with col2:
        st.markdown(
            f"""<div class="metric-container">
                <h3>🏷️ Уникальных товаров</h3>
                <h2>{df['Art'].nunique():,}</h2>
            </div>""", 
            unsafe_allow_html=True
        )
        
    with col3:
        st.markdown(
            f"""<div class="metric-container">
                <h3>🏪 Магазинов</h3>
                <h2>{df['Magazin'].nunique()}</h2>
            </div>""", 
            unsafe_allow_html=True
        )
        
    with col4:
        st.markdown(
            f"""<div class="metric-container">
                <h3>📂 Сегментов</h3>
                <h2>{df['Segment'].nunique()}</h2>
            </div>""", 
            unsafe_allow_html=True
        )
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.info(f"📅 **Период данных**: {df['Datasales'].min().date()} - {df['Datasales'].max().date()}")
    with col2:
        st.info(f"💰 **Общая выручка**: {df['Sum'].sum():.0f} ГРН")
    with col3:
        st.info(f"📈 **Средние продажи/день**: {df.groupby('Datasales')['Qty'].sum().mean():.1f} шт.")

def remove_outliers_iqr(data, multiplier=1.5):
    """ИСПРАВЛЕНО: Удаляет выбросы методом IQR с корректным расчетом границ"""
    if len(data) < 4:
        return data
    
    Q1 = data.quantile(0.25)
    Q3 = data.quantile(0.75)
    IQR = Q3 - Q1
    
    # ИСПРАВЛЕНИЕ: правильный расчет границ
    lower_bound = Q1 - multiplier * IQR
    upper_bound = Q3 + multiplier * IQR
    
    return data.clip(lower=lower_bound, upper=upper_bound)

def smooth_data(data, method='ma', window=7):
    """Сглаживает данные различными методами"""
    if method == 'ma':
        return data.rolling(window=window, min_periods=1, center=True).mean()
    elif method == 'ema':
        return data.ewm(span=window, adjust=False).mean()
    elif method == 'savgol' and len(data) >= window:
        # ИСПРАВЛЕНИЕ: проверка четности окна для Savitzky-Golay
        if window % 2 == 0:
            window += 1
        try:
            return pd.Series(savgol_filter(data, window_length=window, polyorder=min(3, window-1)), index=data.index)
        except:
            return data.rolling(window=window, min_periods=1, center=True).mean()
    else:
        return data

def prepare_prophet_data(df, remove_outliers=False, smooth_method=None, smooth_window=7):
    """ИСПРАВЛЕНО: Подготавливает данные для Prophet с корректной агрегацией"""
    # ИСПРАВЛЕНИЕ: агрегация данных по дате без суммы, только количество
    daily_sales = df.groupby('Datasales')['Qty'].sum().reset_index()
    daily_sales.columns = ['ds', 'y']
    
    original_data = daily_sales.copy()
    
    # Обработка выбросов
    if remove_outliers:
        daily_sales['y'] = remove_outliers_iqr(daily_sales['y'])
    
    # Сглаживание
    if smooth_method:
        daily_sales['y'] = smooth_data(daily_sales['y'], method=smooth_method, window=smooth_window)
    
    # ИСПРАВЛЕНИЕ: заменяем отрицательные значения на 0 вместо NaN
    daily_sales['y'] = daily_sales['y'].clip(lower=0)
    
    return daily_sales, original_data

def train_prophet_model(data, periods=30):
    """Обучает модель Prophet"""
    try:
        model = Prophet(
            daily_seasonality=False,
            weekly_seasonality=True,
            yearly_seasonality=True,
            seasonality_mode='multiplicative',
            changepoint_prior_scale=0.05,
            seasonality_prior_scale=10
        )
        
        model.fit(data)
        
        future = model.make_future_dataframe(periods=periods)
        forecast = model.predict(future)
        
        # ИСПРАВЛЕНИЕ: обеспечиваем неотрицательные прогнозы
        forecast['yhat'] = forecast['yhat'].clip(lower=0)
        forecast['yhat_lower'] = forecast['yhat_lower'].clip(lower=0)
        forecast['yhat_upper'] = forecast['yhat_upper'].clip(lower=0)
        
        return model, forecast
        
    except Exception as e:
        st.error(f"❌ Ошибка при обучении модели: {str(e)}")
        return None, None

def calculate_model_accuracy(train_data, model):
    """ИСПРАВЛЕНО: Корректный расчет метрик точности"""
    try:
        # Прогноз на исторических данных
        historical_forecast = model.predict(train_data[['ds']])
        
        y_true = train_data['y'].values
        y_pred = historical_forecast['yhat'].values
        
        # ИСПРАВЛЕНИЕ: обеспечиваем одинаковую длину массивов
        min_len = min(len(y_true), len(y_pred))
        y_true = y_true[:min_len]
        y_pred = y_pred[:min_len]
        
        # Расчет метрик
        mae = mean_absolute_error(y_true, y_pred)
        rmse = np.sqrt(mean_squared_error(y_true, y_pred))
        
        # ИСПРАВЛЕНИЕ: безопасный расчет MAPE
        mask = y_true != 0
        if mask.sum() > 0:
            mape = np.mean(np.abs((y_true[mask] - y_pred[mask]) / y_true[mask])) * 100
        else:
            mape = 0
        
        # R²
        r2 = r2_score(y_true, y_pred)
        
        return {
            'MAE': mae,
            'RMSE': rmse,
            'MAPE': mape,
            'R2': r2
        }
    except Exception as e:
        st.warning(f"Не удалось рассчитать метрики точности: {str(e)}")
        return None

def show_accuracy_table(metrics):
    """Отображает таблицу метрик точности"""
    st.markdown('<div class="accuracy-card">', unsafe_allow_html=True)
    st.markdown("### 🎯 Метрики точности модели")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("MAE", f"{metrics['MAE']:.2f}")
    with col2:
        st.metric("RMSE", f"{metrics['RMSE']:.2f}")
    with col3:
        st.metric("MAPE", f"{metrics['MAPE']:.2f}%")
    with col4:
        st.metric("R²", f"{metrics['R2']:.4f}")
    
    st.markdown('</div>', unsafe_allow_html=True)

def plot_data_preprocessing(original, processed, title):
    """Визуализирует эффект предобработки данных"""
    fig = go.Figure()
    
    fig.add_trace(go.Scatter(
        x=original['ds'], 
        y=original['y'],
        mode='lines',
        name='Оригинальные данные',
        line=dict(color='lightgray', width=1),
        opacity=0.5
    ))
    
    fig.add_trace(go.Scatter(
        x=processed['ds'],
        y=processed['y'],
        mode='lines',
        name='Обработанные данные',
        line=dict(color='#667eea', width=2)
    ))
    
    fig.update_layout(
        title=title,
        xaxis_title="Дата",
        yaxis_title="Количество",
        hovermode='x unified',
        height=400
    )
    
    return fig

def plot_forecast(train_data, forecast, title):
    """Визуализирует прогноз"""
    fig = go.Figure()
    
    # Исторические данные
    fig.add_trace(go.Scatter(
        x=train_data['ds'],
        y=train_data['y'],
        mode='lines',
        name='Фактические продажи',
        line=dict(color='#1f77b4', width=2)
    ))
    
    # Прогноз
    forecast_future = forecast[forecast['ds'] > train_data['ds'].max()]
    
    fig.add_trace(go.Scatter(
        x=forecast_future['ds'],
        y=forecast_future['yhat'],
        mode='lines',
        name='Прогноз',
        line=dict(color='#ff7f0e', width=2, dash='dash')
    ))
    
    # Доверительный интервал
    fig.add_trace(go.Scatter(
        x=forecast_future['ds'].tolist() + forecast_future['ds'].tolist()[::-1],
        y=forecast_future['yhat_upper'].tolist() + forecast_future['yhat_lower'].tolist()[::-1],
        fill='toself',
        fillcolor='rgba(255, 127, 14, 0.2)',
        line=dict(color='rgba(255, 127, 14, 0)'),
        name='Доверительный интервал',
        showlegend=True
    ))
    
    fig.update_layout(
        title=title,
        xaxis_title="Дата",
        yaxis_title="Количество",
        hovermode='x unified',
        height=500
    )
    
    return fig

def plot_prophet_components(model, forecast):
    """Визуализирует компоненты модели Prophet"""
    fig = go.Figure()
    
    # Тренд
    fig.add_trace(go.Scatter(
        x=forecast['ds'],
        y=forecast['trend'],
        mode='lines',
        name='Тренд',
        line=dict(color='#2ca02c', width=2)
    ))
    
    fig.update_layout(
        title="📊 Декомпозиция: Тренд",
        xaxis_title="Дата",
        yaxis_title="Значение тренда",
        hovermode='x unified',
        height=400
    )
    
    return fig

def calculate_segment_volatility(df, magazin, segment):
    """ИСПРАВЛЕНО: Корректный расчет волатильности сегмента"""
    filtered = df[(df['Magazin'] == magazin) & (df['Segment'] == segment)]
    
    if len(filtered) < 2:
        return 0.3  # Значение по умолчанию
    
    daily_sales = filtered.groupby('Datasales')['Qty'].sum()
    
    if daily_sales.mean() == 0:
        return 0.3
    
    # ИСПРАВЛЕНИЕ: нормализованная волатильность (коэффициент вариации)
    volatility = daily_sales.std() / daily_sales.mean()
    
    # Ограничиваем значение от 0 до 1
    return min(max(volatility, 0), 1)

def get_forecast_scenarios(forecast_df, volatility):
    """ИСПРАВЛЕНО: Корректный расчет сценариев прогноза"""
    realistic = forecast_df['yhat'].values
    
    # ИСПРАВЛЕНИЕ: используем доверительные интервалы Prophet
    lower_bound = forecast_df['yhat_lower'].values
    upper_bound = forecast_df['yhat_upper'].values
    
    # Пессимистичный и оптимистичный сценарии
    pessimistic = np.maximum(lower_bound, 0)
    optimistic = np.maximum(upper_bound, 0)
    
    return realistic, optimistic, pessimistic

def show_forecast_statistics(filtered_df, forecast, forecast_days, magazin, segment, full_df):
    """Показывает статистику прогноза"""
    st.markdown("## 📊 Статистика прогноза")
    
    future_forecast = forecast.tail(forecast_days)
    avg_forecast = future_forecast['yhat'].mean()
    total_forecast = future_forecast['yhat'].sum()
    
    # ИСПРАВЛЕНИЕ: безопасный расчет средней цены
    if len(filtered_df) > 0 and filtered_df['Qty'].sum() > 0:
        avg_price = filtered_df['Sum'].sum() / filtered_df['Qty'].sum()
    else:
        avg_price = 0
    
    forecast_revenue = total_forecast * avg_price
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.metric(
            "📦 Прогноз (средний/день)",
            f"{avg_forecast:.0f} шт."
        )
    
    with col2:
        st.metric(
            f"📊 Прогноз на {forecast_days} дней",
            f"{total_forecast:.0f} шт."
        )
    
    with col3:
        st.metric(
            "💰 Прогноз выручки",
            f"{forecast_revenue:.0f} ГРН"
        )

def plot_monthly_analysis_with_forecast(df, magazin, segment, model, forecast_days, remove_outliers, smooth_method):
    """Расширенный анализ по месяцам с множественными графиками и статистикой"""
    # Фильтрация данных
    filtered = df.copy()
    
    if magazin != 'Все магазины':
        filtered = filtered[filtered['Magazin'] == magazin]
    
    if segment != 'Все сегменты':
        filtered = filtered[filtered['Segment'] == segment]
    
    if len(filtered) == 0:
        st.warning("⚠️ Недостаточно данных для месячного анализа")
        return
    
    # Группировка по месяцам
    filtered['Month'] = filtered['Datasales'].dt.to_period('M')
    monthly_data = filtered.groupby('Month').agg({
        'Qty': 'sum',
        'Sum': 'sum',
        'Art': 'nunique'
    }).reset_index()
    monthly_data.columns = ['Month', 'Qty', 'Sum', 'Unique_Products']
    
    monthly_data['Month'] = monthly_data['Month'].dt.to_timestamp()
    
    # Расчет средней цены и других метрик
    monthly_data['Avg_Price'] = np.where(
        monthly_data['Qty'] > 0,
        monthly_data['Sum'] / monthly_data['Qty'],
        0
    )
    
    # Прогноз на будущий месяц
    future_dates = pd.date_range(
        start=filtered['Datasales'].max() + pd.Timedelta(days=1),
        periods=forecast_days
    )
    future_df = pd.DataFrame({'ds': future_dates})
    future_forecast = model.predict(future_df)
    
    # Агрегация прогноза по месяцам
    future_forecast['Month'] = pd.to_datetime(future_forecast['ds']).dt.to_period('M').dt.to_timestamp()
    forecast_monthly = future_forecast.groupby('Month')['yhat'].sum().reset_index()
    forecast_monthly.columns = ['Month', 'Forecast_Qty']
    
    last_avg_price = monthly_data['Avg_Price'].iloc[-1] if len(monthly_data) > 0 else 0
    forecast_monthly['Forecast_Revenue'] = forecast_monthly['Forecast_Qty'] * last_avg_price
    
    # === СТАТИСТИКА МЕСЯЧНОГО АНАЛИЗА ===
    st.markdown("### 📊 Статистика по месяцам")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        avg_monthly_sales = monthly_data['Qty'].mean()
        st.metric(
            "📦 Средние продажи/месяц",
            f"{avg_monthly_sales:.0f} шт",
            delta=f"{monthly_data['Qty'].iloc[-1] - avg_monthly_sales:.0f}"
        )
    
    with col2:
        avg_monthly_revenue = monthly_data['Sum'].mean()
        st.metric(
            "💰 Средняя выручка/месяц",
            f"{avg_monthly_revenue:.0f} ГРН",
            delta=f"{monthly_data['Sum'].iloc[-1] - avg_monthly_revenue:.0f}"
        )
    
    with col3:
        # Рост продаж (последний месяц vs предыдущий)
        if len(monthly_data) >= 2:
            growth_rate = ((monthly_data['Qty'].iloc[-1] / monthly_data['Qty'].iloc[-2]) - 1) * 100
            st.metric(
                "📈 Рост (месяц к месяцу)",
                f"{growth_rate:+.1f}%",
                delta=f"{monthly_data['Qty'].iloc[-1] - monthly_data['Qty'].iloc[-2]:.0f} шт"
            )
        else:
            st.metric("📈 Рост", "N/A")
    
    with col4:
        # Прогноз роста
        if len(forecast_monthly) > 0 and len(monthly_data) > 0:
            # Сравниваем последний прогнозный месяц с последним фактическим месяцем
            forecast_growth = ((forecast_monthly['Forecast_Qty'].iloc[-1] / monthly_data['Qty'].iloc[-1]) - 1) * 100
            st.metric(
                "🔮 Прогноз роста",
                f"{forecast_growth:+.1f}%",
                delta="прогнозный период"
            )
        else:
            st.metric("🔮 Прогноз роста", "N/A")
    
    # === ГРАФИК 1: Продажи и выручка ===
    st.markdown("### 📊 График продаж и выручки")
    
    fig1 = go.Figure()
    
    # Фактические продажи
    fig1.add_trace(go.Bar(
        x=monthly_data['Month'],
        y=monthly_data['Qty'],
        name='Фактические продажи',
        marker_color='#1f77b4',
        yaxis='y',
        text=monthly_data['Qty'].round(0),
        textposition='outside',
        texttemplate='%{text:.0f}'
    ))
    
    # Прогнозные продажи
    fig1.add_trace(go.Bar(
        x=forecast_monthly['Month'],
        y=forecast_monthly['Forecast_Qty'],
        name='Прогноз продаж',
        marker_color='#ff7f0e',
        yaxis='y',
        text=forecast_monthly['Forecast_Qty'].round(0),
        textposition='outside',
        texttemplate='%{text:.0f}'
    ))
    
    # Фактическая выручка
    fig1.add_trace(go.Scatter(
        x=monthly_data['Month'],
        y=monthly_data['Sum'],
        name='Фактическая выручка',
        mode='lines+markers',
        line=dict(color='#2ca02c', width=3),
        marker=dict(size=8),
        yaxis='y2'
    ))
    
    # Прогнозная выручка
    fig1.add_trace(go.Scatter(
        x=forecast_monthly['Month'],
        y=forecast_monthly['Forecast_Revenue'],
        name='Прогноз выручки',
        mode='lines+markers',
        line=dict(color='#d62728', width=3, dash='dash'),
        marker=dict(size=8, symbol='diamond'),
        yaxis='y2'
    ))
    
    fig1.update_layout(
        xaxis_title="Месяц",
        yaxis=dict(title="Количество (шт)", side='left'),
        yaxis2=dict(title="Выручка (ГРН)", side='right', overlaying='y'),
        hovermode='x unified',
        height=500,
        legend=dict(x=0, y=1.15, orientation='h'),
        barmode='group'
    )
    
    st.plotly_chart(fig1, use_container_width=True, key="plot_forecast_fig1")
    
    # === ГРАФИК 2: Динамика средней цены ===
    st.markdown("### 💵 Динамика средней цены")
    
    fig2 = go.Figure()
    
    fig2.add_trace(go.Scatter(
        x=monthly_data['Month'],
        y=monthly_data['Avg_Price'],
        mode='lines+markers',
        name='Средняя цена',
        line=dict(color='#9467bd', width=3),
        marker=dict(size=10),
        fill='tozeroy',
        fillcolor='rgba(148, 103, 189, 0.2)'
    ))
    
    # Добавляем среднюю линию
    avg_price_overall = monthly_data['Avg_Price'].mean()
    fig2.add_hline(
        y=avg_price_overall,
        line_dash="dash",
        line_color="red",
        annotation_text=f"Среднее: {avg_price_overall:.2f} ГРН",
        annotation_position="right"
    )
    
    fig2.update_layout(
        xaxis_title="Месяц",
        yaxis_title="Средняя цена (ГРН)",
        hovermode='x unified',
        height=400
    )
    
    st.plotly_chart(fig2, use_container_width=True, key="plot_forecast_fig2")
    
    # === ГРАФИК 3: Уникальные товары и интенсивность продаж ===
    st.markdown("### 🏷️ Ассортимент и интенсивность продаж")
    
    col1, col2 = st.columns(2)
    
    with col1:
        fig3 = go.Figure()
        
        fig3.add_trace(go.Bar(
            x=monthly_data['Month'],
            y=monthly_data['Unique_Products'],
            name='Уникальных товаров',
            marker_color='#8c564b',
            text=monthly_data['Unique_Products'],
            textposition='outside'
        ))
        
        fig3.update_layout(
            title="Количество уникальных товаров",
            xaxis_title="Месяц",
            yaxis_title="Кол-во товаров",
            height=350
        )
        
        st.plotly_chart(fig3, use_container_width=True, key="plot_forecast_fig3")
    
    with col2:
        # Интенсивность продаж (продажи на 1 товар)
        monthly_data['Sales_per_Product'] = monthly_data['Qty'] / monthly_data['Unique_Products']
        
        fig4 = go.Figure()
        
        fig4.add_trace(go.Scatter(
            x=monthly_data['Month'],
            y=monthly_data['Sales_per_Product'],
            mode='lines+markers',
            name='Продаж на товар',
            line=dict(color='#e377c2', width=3),
            marker=dict(size=10),
            fill='tozeroy',
            fillcolor='rgba(227, 119, 194, 0.2)'
        ))
        
        fig4.update_layout(
            title="Интенсивность продаж (шт/товар)",
            xaxis_title="Месяц",
            yaxis_title="Продаж на 1 товар",
            height=350
        )
        
        st.plotly_chart(fig4, use_container_width=True, key="plot_forecast_fig4")
    
    # === ГРАФИК 4: Сравнение факт vs прогноз ===
    if len(forecast_monthly) > 0:
        st.markdown("### 🔮 Сравнительный анализ: Факт vs Прогноз")
        
        # Берем последние 6 месяцев истории
        recent_months = monthly_data.tail(6)
        
        comparison_data = pd.DataFrame({
            'Категория': ['Последние 6 мес (факт)'] * len(recent_months) + ['Прогноз'] * len(forecast_monthly),
            'Период': list(recent_months['Month']) + list(forecast_monthly['Month']),
            'Продажи': list(recent_months['Qty']) + list(forecast_monthly['Forecast_Qty']),
            'Выручка': list(recent_months['Sum']) + list(forecast_monthly['Forecast_Revenue'])
        })
        
        fig5 = go.Figure()
        
        colors = {'Последние 6 мес (факт)': '#3498db', 'Прогноз': '#e74c3c'}
        
        for category in comparison_data['Категория'].unique():
            cat_data = comparison_data[comparison_data['Категория'] == category]
            
            fig5.add_trace(go.Bar(
                x=cat_data['Период'],
                y=cat_data['Продажи'],
                name=category,
                marker_color=colors[category],
                text=cat_data['Продажи'].round(0),
                textposition='outside'
            ))
        
        fig5.update_layout(
            title="Сравнение факт vs прогноз",
            xaxis_title="Период",
            yaxis_title="Продажи (шт)",
            height=400,
            barmode='group'
        )
        
        st.plotly_chart(fig5, use_container_width=True, key="plot_forecast_fig5")
    
    # === ТАБЛИЦА С ДЕТАЛЬНОЙ СТАТИСТИКОЙ ===
    st.markdown("### 📋 Детальная статистика по месяцам")
    
    # Создаем расширенную таблицу
    display_table = monthly_data.copy()
    display_table['Month'] = display_table['Month'].dt.strftime('%Y-%m')
    
    # Добавляем процент изменения
    display_table['Qty_Change_%'] = display_table['Qty'].pct_change() * 100
    display_table['Revenue_Change_%'] = display_table['Sum'].pct_change() * 100
    
    display_table = display_table.rename(columns={
        'Month': '📅 Месяц',
        'Qty': '📦 Продажи',
        'Sum': '💰 Выручка',
        'Avg_Price': '💵 Средняя цена',
        'Unique_Products': '🏷️ Товаров',
        'Sales_per_Product': '📊 Продаж/товар',
        'Qty_Change_%': '📈 Δ Продаж %',
        'Revenue_Change_%': '💹 Δ Выручка %'
    })
    
    # Форматирование
    st.dataframe(
        display_table.style.format({
            '📦 Продажи': '{:.0f}',
            '💰 Выручка': '{:.0f}',
            '💵 Средняя цена': '{:.2f}',
            '🏷️ Товаров': '{:.0f}',
            '📊 Продаж/товар': '{:.1f}',
            '📈 Δ Продаж %': '{:+.1f}%',
            '💹 Δ Выручка %': '{:+.1f}%'
        }).background_gradient(subset=['📈 Δ Продаж %', '💹 Δ Выручка %'], cmap='RdYlGn', vmin=-20, vmax=20),
        use_container_width=True,
        hide_index=True
    )
    
    # === АЛЕРТЫ И РЕКОМЕНДАЦИИ ===
    st.markdown("### 🚨 Алерты и рекомендации")
    
    alerts = []
    recommendations = []
    
    # Проверка падения продаж
    if len(monthly_data) >= 2:
        last_month_sales = monthly_data['Qty'].iloc[-1]
        prev_month_sales = monthly_data['Qty'].iloc[-2]
        
        if last_month_sales < prev_month_sales * 0.8:
            alerts.append("📉 **КРИТИЧНО**: Падение продаж более 20% за последний месяц!")
            recommendations.append("🎯 Срочно запустить промо-акции и проанализировать причины падения")
        elif last_month_sales < prev_month_sales * 0.9:
            alerts.append("⚠️ Снижение продаж на 10-20% за последний месяц")
            recommendations.append("📊 Провести анализ конкурентов и пересмотреть ценовую политику")
    
    # Проверка роста цены
    if len(monthly_data) >= 3:
        recent_price = monthly_data['Avg_Price'].tail(3).mean()
        older_price = monthly_data['Avg_Price'].head(3).mean()
        
        if recent_price > older_price * 1.15:
            alerts.append("💵 Средняя цена выросла более чем на 15%")
            recommendations.append("⚖️ Проверить влияние роста цен на объем продаж")
    
    # Проверка ассортимента
    if len(monthly_data) >= 2:
        last_products = monthly_data['Unique_Products'].iloc[-1]
        avg_products = monthly_data['Unique_Products'].mean()
        
        if last_products < avg_products * 0.7:
            alerts.append("🏷️ Резкое сокращение ассортимента")
            recommendations.append("📦 Расширить ассортимент товаров для увеличения продаж")
    
    # Проверка прогноза
    if len(forecast_monthly) > 0 and len(monthly_data) > 0:
        forecast_vs_last = forecast_monthly['Forecast_Qty'].iloc[0] / monthly_data['Qty'].iloc[-1]
        
        if forecast_vs_last > 1.3:
            alerts.append("🚀 Прогноз показывает рост продаж более 30%")
            recommendations.append("📦 Увеличить закупки и подготовить склад к повышенному спросу")
        elif forecast_vs_last < 0.8:
            alerts.append("📉 Прогноз указывает на снижение продаж")
            recommendations.append("💡 Рассмотреть промо-кампании для стимулирования спроса")
    
    # Сезонность
    if len(monthly_data) >= 6:
        monthly_data_temp = monthly_data.copy()
        monthly_data_temp['Month_Num'] = pd.to_datetime(monthly_data_temp['Month']).dt.month
        seasonality = monthly_data_temp.groupby('Month_Num')['Qty'].mean().std()
        
        if seasonality > monthly_data['Qty'].mean() * 0.3:
            recommendations.append("📅 Выявлена высокая сезонность - планируйте закупки с учетом сезонных колебаний")
    
    # Отображение алертов
    if alerts:
        for alert in alerts:
            st.markdown(f'<div class="problem-card">{alert}</div>', unsafe_allow_html=True)
    else:
        st.success("✅ Критических проблем не обнаружено")
    
    # Отображение рекомендаций
    if recommendations:
        st.markdown("#### 💡 Рекомендации:")
        for rec in recommendations:
            st.markdown(f'<div class="insight-card">{rec}</div>', unsafe_allow_html=True)
    
    # Дополнительные рекомендации
    st.markdown("#### 🎯 Общие рекомендации:")
    
    general_recommendations = [
        f"📊 Средняя цена: {monthly_data['Avg_Price'].mean():.2f} ГРН - {'оптимальна' if monthly_data['Avg_Price'].std() < monthly_data['Avg_Price'].mean() * 0.2 else 'сильно варьируется'}",
        f"📦 Оптимальный запас на месяц: {monthly_data['Qty'].mean() * 1.2:.0f} единиц (среднее + 20% буфер)",
        f"💰 Целевая выручка на следующий месяц: {monthly_data['Sum'].mean() * 1.1:.0f} ГРН (+10% к среднему)"
    ]
    
    for rec in general_recommendations:
        st.info(rec)
    
def get_top_models_by_segment(df, magazin):
    """Получает топ-10 моделей по каждому сегменту"""
    filtered = df[df['Magazin'] == magazin]
    segments = filtered['Segment'].unique()
    
    result = {}
    
    for segment in segments:
        segment_data = filtered[filtered['Segment'] == segment]
        
        top_models = segment_data.groupby('Model').agg({
            'Qty': 'sum',
            'Sum': 'sum'
        }).reset_index()
        
        # ИСПРАВЛЕНИЕ: безопасный расчет средней цены
        top_models['Price'] = np.where(
            top_models['Qty'] > 0,
            top_models['Sum'] / top_models['Qty'],
            0
        )
        
        top_models = top_models.sort_values('Sum', ascending=False).head(10)
        result[segment] = top_models
    
    return result

def generate_insights(df, forecast, magazin, segment):
    """Генерирует инсайты и рекомендации"""
    filtered = df[(df['Magazin'] == magazin) & (df['Segment'] == segment)]
    
    insights = []
    problems = []
    
    if len(filtered) == 0:
        return insights, problems
    
    # Анализ тренда
    recent_sales = filtered.tail(30)['Qty'].sum()
    older_sales = filtered.head(30)['Qty'].sum()
    
    if recent_sales > older_sales * 1.2:
        insights.append("📈 Продажи растут! Рекомендуется увеличить закупки.")
    elif recent_sales < older_sales * 0.8:
        problems.append("📉 Снижение продаж. Необходим анализ причин.")
        insights.append("🔍 Рассмотрите проведение промо-акций.")
    
    # Анализ волатильности
    daily_sales = filtered.groupby('Datasales')['Qty'].sum()
    cv = daily_sales.std() / daily_sales.mean() if daily_sales.mean() > 0 else 0
    
    if cv > 0.5:
        problems.append("⚠️ Высокая волатильность продаж. Сложно планировать запасы.")
        insights.append("📦 Рекомендуется создать буферный запас.")
    
    # Анализ прогноза
    future_forecast = forecast.tail(30)
    avg_forecast = future_forecast['yhat'].mean()
    historical_avg = daily_sales.tail(30).mean()
    
    if avg_forecast > historical_avg * 1.1:
        insights.append("🚀 Прогноз показывает рост продаж. Подготовьте дополнительные запасы.")
    elif avg_forecast < historical_avg * 0.9:
        insights.append("⚡ Ожидается спад продаж. Оптимизируйте закупки.")
    
    return insights, problems

def create_word_report(detailed_forecast, selected_magazin, selected_segment, forecast_days, 
                      total_forecast, avg_daily_forecast, forecast_revenue, confidence_score,
                      accuracy_metrics, insights, filtered_df, prophet_data):
    """Создает Word отчет с результатами прогнозирования"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO
        
        doc = Document()
        
        # Настройка стилей
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        
        # Заголовок
        title = doc.add_heading('ОТЧЕТ ПО ПРОГНОЗИРОВАНИЮ ПРОДАЖ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Информация о прогнозе
        doc.add_heading('Информация о прогнозе', level=1)
        
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Light Grid Accent 1'
        
        info_data = [
            ('Дата создания:', datetime.now().strftime('%Y-%m-%d %H:%M')),
            ('Магазин:', selected_magazin),
            ('Сегмент:', selected_segment),
            ('Период прогноза:', f'{forecast_days} дней')
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_table.rows[i].cells[0].text = label
            info_table.rows[i].cells[1].text = value
        
        doc.add_paragraph()
        
        # Основные показатели
        doc.add_heading('Основные показатели', level=1)
        
        metrics_table = doc.add_table(rows=5, cols=2)
        metrics_table.style = 'Light Grid Accent 1'
        
        metrics_data = [
            ('Общий прогноз', f'{total_forecast:.0f} единиц'),
            ('Средние продажи/день', f'{avg_daily_forecast:.0f} единиц'),
            ('Прогнозная выручка', f'{forecast_revenue:.0f} ГРН'),
            ('Уверенность прогноза', f'{confidence_score:.0f}%')
        ]
        
        metrics_table.rows[0].cells[0].text = 'Показатель'
        metrics_table.rows[0].cells[1].text = 'Значение'
        
        for i, (label, value) in enumerate(metrics_data, 1):
            metrics_table.rows[i].cells[0].text = label
            metrics_table.rows[i].cells[1].text = value
        
        doc.add_paragraph()
        
        # Метрики точности модели
        if accuracy_metrics:
            doc.add_heading('Метрики точности модели', level=1)
            
            accuracy_table = doc.add_table(rows=5, cols=3)
            accuracy_table.style = 'Light Grid Accent 1'
            
            accuracy_table.rows[0].cells[0].text = 'Метрика'
            accuracy_table.rows[0].cells[1].text = 'Значение'
            accuracy_table.rows[0].cells[2].text = 'Интерпретация'
            
            accuracy_data = [
                ('MAE', f"{accuracy_metrics['MAE']:.2f}", 'Средняя абсолютная ошибка'),
                ('RMSE', f"{accuracy_metrics['RMSE']:.2f}", 'Корень из средней квадратичной ошибки'),
                ('MAPE', f"{accuracy_metrics['MAPE']:.2f}%", 'Средняя абсолютная процентная ошибка'),
                ('R²', f"{accuracy_metrics['R2']:.4f}", 'Коэффициент детерминации')
            ]
            
            for i, (metric, value, interpretation) in enumerate(accuracy_data, 1):
                accuracy_table.rows[i].cells[0].text = metric
                accuracy_table.rows[i].cells[1].text = value
                accuracy_table.rows[i].cells[2].text = interpretation
            
            doc.add_paragraph()
        
        # Детальный прогноз
        doc.add_heading('Детальный прогноз (первые 10 дней)', level=1)
        
        forecast_table = doc.add_table(rows=11, cols=4)
        forecast_table.style = 'Light Grid Accent 1'
        
        forecast_table.rows[0].cells[0].text = 'Дата'
        forecast_table.rows[0].cells[1].text = 'Пессимистичный'
        forecast_table.rows[0].cells[2].text = 'Реальный'
        forecast_table.rows[0].cells[3].text = 'Оптимистичный'
        
        for i, (idx, row) in enumerate(detailed_forecast.head(10).iterrows(), 1):
            forecast_table.rows[i].cells[0].text = row['📅 Дата']
            forecast_table.rows[i].cells[1].text = str(row['😰 Пессимистичный'])
            forecast_table.rows[i].cells[2].text = str(row['🎯 Реальный'])
            forecast_table.rows[i].cells[3].text = str(row['🚀 Оптимистичный'])
        
        doc.add_page_break()
        
        # Статистика исторических данных
        doc.add_heading('Статистика исторических данных', level=1)
        
        total_sales = filtered_df['Qty'].sum()
        total_revenue = filtered_df['Sum'].sum()
        avg_price = total_revenue / total_sales if total_sales > 0 else 0
        period_days = (filtered_df['Datasales'].max() - filtered_df['Datasales'].min()).days + 1
        
        stats_table = doc.add_table(rows=10, cols=2)
        stats_table.style = 'Light Grid Accent 1'
        
        stats_data = [
            ('Показатель', 'Значение'),
            ('Период данных', f"{filtered_df['Datasales'].min().strftime('%Y-%m-%d')} - {filtered_df['Datasales'].max().strftime('%Y-%m-%d')}"),
            ('Всего дней', f'{period_days}'),
            ('Всего продано', f'{total_sales:.0f} единиц'),
            ('Общая выручка', f'{total_revenue:.0f} ГРН'),
            ('Средняя цена', f'{avg_price:.2f} ГРН'),
            ('Средние продажи/день', f'{prophet_data["y"].mean():.1f} единиц'),
            ('Макс. продажи за день', f'{prophet_data["y"].max():.0f} единиц'),
            ('Мин. продажи за день', f'{prophet_data["y"].min():.0f} единиц')
        ]
        
        for i, (label, value) in enumerate(stats_data):
            stats_table.rows[i].cells[0].text = label
            stats_table.rows[i].cells[1].text = value
        
        doc.add_paragraph()
        
        # Рекомендации
        doc.add_heading('Рекомендации', level=1)
        
        for i, insight in enumerate(insights[:15], 1):
            # Удаляем эмодзи
            clean_insight = insight.replace('🎯', '').replace('📊', '').replace('🚀', '').replace('⚡', '').replace('📦', '').replace('🔍', '').replace('📈', '').replace('💡', '').replace('🔥', '').strip()
            p = doc.add_paragraph(f'{i}. {clean_insight}')
            p.style = 'List Number'
        
        doc.add_paragraph()
        
        # Заключение
        doc.add_heading('Заключение', level=1)
        
        conclusion_text = f"""
Данный прогноз основан на анализе исторических данных продаж за период с {filtered_df['Datasales'].min().strftime('%Y-%m-%d')} по {filtered_df['Datasales'].max().strftime('%Y-%m-%d')}.

Модель показывает уверенность прогноза на уровне {confidence_score:.0f}%, что {"говорит о высокой надежности" if confidence_score > 70 else "требует осторожного применения"} результатов для планирования.

Рекомендуется регулярно обновлять прогнозы с появлением новых данных для повышения точности планирования.
        """
        
        doc.add_paragraph(conclusion_text.strip())
        
        # Футер
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.add_run(f'Отчет сгенерирован: {datetime.now().strftime("%Y-%m-%d %H:%M")}').italic = True
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Сохранение в BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
        
    except Exception as e:
        st.error(f"Ошибка при создании Word документа: {str(e)}")
        st.info("Установите библиотеку: pip install python-docx")
        return None

def main():
    st.markdown('<h1 class="main-header">🏪 Система прогнозирования продаж</h1>', unsafe_allow_html=True)
    
    with st.sidebar:
        st.markdown("## ⚙️ Настройки")
        
        uploaded_file = st.file_uploader(
            "📁 Загрузите Excel файл",
            type=['xlsx', 'xls'],
            help="Файл должен содержать колонки: Magazin, Datasales, Art, Describe, Model, Segment, Price, Qty, Sum"
        )
        
        st.markdown("---")
        st.markdown("### 🔧 Параметры прогноза")
        
        forecast_days = st.slider(
            "📅 Период прогноза (дней)",
            min_value=7,
            max_value=90,
            value=30,
            step=1
        )
        
        st.markdown("### 🧹 Предобработка данных")
        
        remove_outliers = st.checkbox(
            "Удалить выбросы",
            value=True,
            help="Использует метод IQR для удаления аномальных значений"
        )
        
        smooth_method = st.selectbox(
            "Метод сглаживания",
            options=['none', 'ma', 'ema', 'savgol'],
            format_func=lambda x: {
                'none': 'Без сглаживания',
                'ma': 'Скользящее среднее',
                'ema': 'Экспоненциальное сглаживание',
                'savgol': 'Фильтр Савицкого-Голея'
            }[x]
        )
        
        if smooth_method != 'none':
            smooth_window = st.slider(
                "Окно сглаживания",
                min_value=3,
                max_value=21,
                value=7,
                step=2
            )
        else:
            smooth_window = 7
    
    if uploaded_file is None:
        st.info("👈 Загрузите Excel файл для начала работы")
        
        st.markdown("### 📋 Требования к данным")
        st.markdown("""
        Файл должен содержать следующие колонки:
        - **Magazin**: Название магазина
        - **Datasales**: Дата продажи
        - **Art**: Артикул товара
        - **Describe**: Описание товара
        - **Model**: Модель товара
        - **Segment**: Сегмент товара
        - **Price**: Цена
        - **Qty**: Количество
        - **Sum**: Сумма продажи
        """)
        return
    
    df = load_and_validate_data(uploaded_file)
    
    if df is None:
        return
    
    show_data_statistics(df)
    
    st.markdown("---")
    st.markdown("## 🎯 Выбор параметров анализа")
    
    col1, col2 = st.columns(2)
    
    with col1:
        available_magazins = ['Все магазины'] + sorted(df['Magazin'].unique().tolist())
        selected_magazin = st.selectbox("🏪 Выберите магазин", available_magazins)
    
    with col2:
        if selected_magazin == 'Все магазины':
            available_segments = ['Все сегменты'] + sorted(df['Segment'].unique().tolist())
        else:
            magazin_df = df[df['Magazin'] == selected_magazin]
            available_segments = ['Все сегменты'] + sorted(magazin_df['Segment'].unique().tolist())
        
        selected_segment = st.selectbox("📂 Выберите сегмент", available_segments)
    
    if st.button("🚀 Создать прогноз", type="primary", use_container_width=True):
        with st.spinner("🔄 Обучение модели..."):
            filtered_df = df.copy()
            
            if selected_magazin != 'Все магазины':
                filtered_df = filtered_df[filtered_df['Magazin'] == selected_magazin]
            
            if selected_segment != 'Все сегменты':
                filtered_df = filtered_df[filtered_df['Segment'] == selected_segment]
            
            if len(filtered_df) < 10:
                st.error("❌ Недостаточно данных для прогнозирования (минимум 10 записей)")
                return
            
            prophet_data, original_data = prepare_prophet_data(
                filtered_df, 
                remove_outliers=remove_outliers, 
                smooth_method=smooth_method if smooth_method != 'none' else None,
                smooth_window=smooth_window
            )
            
            if remove_outliers or (smooth_method and smooth_method != 'none'):
                st.markdown("## 🧹 Предварительная обработка данных")
                
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown("### 📊 Статистика до обработки")
                    st.metric("Среднее", f"{original_data['y'].mean():.2f}")
                    st.metric("Std. отклонение", f"{original_data['y'].std():.2f}")
                    volatility_before = (original_data['y'].std()/original_data['y'].mean()*100) if original_data['y'].mean() > 0 else 0
                    st.metric("Волатильность", f"{volatility_before:.1f}%")
                
                with col2:
                    st.markdown("### ✨ Статистика после обработки")
                    st.metric("Среднее", f"{prophet_data['y'].mean():.2f}", 
                             delta=f"{prophet_data['y'].mean() - original_data['y'].mean():.2f}")
                    st.metric("Std. отклонение", f"{prophet_data['y'].std():.2f}", 
                             delta=f"{prophet_data['y'].std() - original_data['y'].std():.2f}")
                    volatility_after = (prophet_data['y'].std()/prophet_data['y'].mean()*100) if prophet_data['y'].mean() > 0 else 0
                    st.metric("Волатильность", f"{volatility_after:.1f}%", 
                             delta=f"{volatility_after - volatility_before:.1f}%")
                
                fig_preprocessing = plot_data_preprocessing(original_data, prophet_data, "🔄 Сравнение: Оригинальные vs Обработанные данные")
                st.plotly_chart(fig_preprocessing, use_container_width=True, key="preprocessing")
            
            model, forecast = train_prophet_model(prophet_data, periods=forecast_days)
            
            if model is None or forecast is None:
                return
            
            st.success("✅ Модель успешно обучена!")
            
            accuracy_metrics = calculate_model_accuracy(prophet_data, model)
            if accuracy_metrics:
                show_accuracy_table(accuracy_metrics)
            
            show_forecast_statistics(filtered_df, forecast, forecast_days, selected_magazin, selected_segment, df)
            
            st.markdown("## 📈 Прогноз продаж")
            
            fig_main = plot_forecast(
                prophet_data, 
                forecast, 
                f"Прогноз продаж - {selected_magazin} / {selected_segment}"
            )
            st.plotly_chart(fig_main, use_container_width=True, key="main_forecast")
            
            # Добавляем месячный анализ сразу после прогноза
            st.markdown("## 📊 Анализ по месяцам с прогнозом выручки")
            plot_monthly_analysis_with_forecast(
                df, selected_magazin, selected_segment, model, 
                forecast_days, remove_outliers, smooth_method if smooth_method != 'none' else None
            )
            
            st.markdown("## 🔍 Детальный анализ")
            
            fig_components = plot_prophet_components(model, forecast)
            st.plotly_chart(fig_components, use_container_width=True, key="prophet_components")
            
            # === АНАЛИЗ ДНЯ НЕДЕЛИ ===
            st.markdown("### 📅 Анализ продаж по дням недели")
            
            # Подготовка данных по дням недели
            filtered_df_weekday = filtered_df.copy()
            filtered_df_weekday['Weekday'] = filtered_df_weekday['Datasales'].dt.dayofweek
            filtered_df_weekday['Weekday_Name'] = filtered_df_weekday['Datasales'].dt.day_name()
            
            # Переводим названия дней
            weekday_translation = {
                'Monday': 'Понедельник',
                'Tuesday': 'Вторник',
                'Wednesday': 'Среда',
                'Thursday': 'Четверг',
                'Friday': 'Пятница',
                'Saturday': 'Суббота',
                'Sunday': 'Воскресенье'
            }
            filtered_df_weekday['Weekday_Name_RU'] = filtered_df_weekday['Weekday_Name'].map(weekday_translation)
            
            weekday_stats = filtered_df_weekday.groupby(['Weekday', 'Weekday_Name_RU']).agg({
                'Qty': 'sum',
                'Sum': 'sum'
            }).reset_index().sort_values('Weekday')
            
            weekday_stats['Avg_Price'] = weekday_stats['Sum'] / weekday_stats['Qty']
            weekday_stats['Qty_Percent'] = (weekday_stats['Qty'] / weekday_stats['Qty'].sum() * 100)
            
            col1, col2 = st.columns(2)
            
            with col1:
                # График продаж по дням недели
                fig_weekday1 = go.Figure()
                
                colors = ['#ff6b6b' if qty == weekday_stats['Qty'].max() else '#1f77b4' 
                         for qty in weekday_stats['Qty']]
                
                fig_weekday1.add_trace(go.Bar(
                    x=weekday_stats['Weekday_Name_RU'],
                    y=weekday_stats['Qty'],
                    marker_color=colors,
                    text=weekday_stats['Qty'].round(0),
                    textposition='outside',
                    texttemplate='%{text:.0f}<br>(%{customdata:.1f}%)',
                    customdata=weekday_stats['Qty_Percent'],
                    name='Продажи'
                ))
                
                fig_weekday1.update_layout(
                    title="📦 Продажи по дням недели",
                    xaxis_title="День недели",
                    yaxis_title="Количество продаж",
                    height=400,
                    showlegend=False
                )
                
                st.plotly_chart(fig_weekday1, use_container_width=True, key="weekday_sales")
                
                # Находим лучший и худший день
                best_day = weekday_stats.loc[weekday_stats['Qty'].idxmax(), 'Weekday_Name_RU']
                worst_day = weekday_stats.loc[weekday_stats['Qty'].idxmin(), 'Weekday_Name_RU']
                best_qty = weekday_stats['Qty'].max()
                worst_qty = weekday_stats['Qty'].min()
                
                st.success(f"🏆 **Лучший день**: {best_day} ({best_qty:.0f} шт)")
                st.error(f"📉 **Слабый день**: {worst_day} ({worst_qty:.0f} шт)")
            
            with col2:
                # График выручки по дням недели
                fig_weekday2 = go.Figure()
                
                colors_revenue = ['#2ecc71' if rev == weekday_stats['Sum'].max() else '#3498db' 
                                 for rev in weekday_stats['Sum']]
                
                fig_weekday2.add_trace(go.Bar(
                    x=weekday_stats['Weekday_Name_RU'],
                    y=weekday_stats['Sum'],
                    marker_color=colors_revenue,
                    text=weekday_stats['Sum'].round(0),
                    textposition='outside',
                    texttemplate='%{text:.0f}',
                    name='Выручка'
                ))
                
                fig_weekday2.update_layout(
                    title="💰 Выручка по дням недели",
                    xaxis_title="День недели",
                    yaxis_title="Выручка (ГРН)",
                    height=400,
                    showlegend=False
                )
                
                st.plotly_chart(fig_weekday2, use_container_width=True, key="weekday_revenue")
                
                best_revenue_day = weekday_stats.loc[weekday_stats['Sum'].idxmax(), 'Weekday_Name_RU']
                best_revenue = weekday_stats['Sum'].max()
                
                st.success(f"💎 **Максимальная выручка**: {best_revenue_day} ({best_revenue:.0f} ГРН)")
            
            # Круговая диаграмма распределения продаж
            st.markdown("### 🥧 Распределение продаж по дням недели")
            
            fig_pie = go.Figure(data=[go.Pie(
                labels=weekday_stats['Weekday_Name_RU'],
                values=weekday_stats['Qty'],
                hole=.4,
                marker=dict(colors=['#FF6384', '#36A2EB', '#FFCE56', '#4BC0C0', '#9966FF', '#FF9F40', '#FF6384']),
                textinfo='label+percent',
                textposition='outside'
            )])
            
            fig_pie.update_layout(
                title="Доля продаж по дням недели",
                height=450,
                showlegend=True
            )
            
            st.plotly_chart(fig_pie, use_container_width=True, key="weekday_pie")
            
            # Рекомендации по дням недели
            st.markdown("#### 💡 Рекомендации по дням недели:")
            
            avg_qty = weekday_stats['Qty'].mean()
            weak_days = weekday_stats[weekday_stats['Qty'] < avg_qty * 0.8]['Weekday_Name_RU'].tolist()
            strong_days = weekday_stats[weekday_stats['Qty'] > avg_qty * 1.2]['Weekday_Name_RU'].tolist()
            
            if weak_days:
                st.markdown(f'<div class="problem-card">📉 Слабые дни ({", ".join(weak_days)}): Проведите акции или скидки для стимулирования продаж</div>', unsafe_allow_html=True)
            
            if strong_days:
                st.markdown(f'<div class="insight-card">🚀 Сильные дни ({", ".join(strong_days)}): Обеспечьте достаточный запас товаров и персонала</div>', unsafe_allow_html=True)
            
            # === ДОПОЛНИТЕЛЬНЫЕ ГРАФИКИ ===
            st.markdown("### 📊 Дополнительная аналитика")
            
            # Тепловая карта продаж: день недели x неделя месяца
            filtered_df_heatmap = filtered_df.copy()
            filtered_df_heatmap['Week'] = filtered_df_heatmap['Datasales'].dt.isocalendar().week
            filtered_df_heatmap['Weekday'] = filtered_df_heatmap['Datasales'].dt.dayofweek
            
            heatmap_data = filtered_df_heatmap.groupby(['Week', 'Weekday'])['Qty'].sum().reset_index()
            heatmap_pivot = heatmap_data.pivot(index='Week', columns='Weekday', values='Qty').fillna(0)
            
            # Названия дней для колонок
            day_names = ['Пн', 'Вт', 'Ср', 'Чт', 'Пт', 'Сб', 'Вс']
            heatmap_pivot.columns = [day_names[i] if i < len(day_names) else str(i) for i in heatmap_pivot.columns]
            
            fig_heatmap = go.Figure(data=go.Heatmap(
                z=heatmap_pivot.values,
                x=heatmap_pivot.columns,
                y=[f'Неделя {int(w)}' for w in heatmap_pivot.index],
                colorscale='Blues',
                text=heatmap_pivot.values.round(0),
                texttemplate='%{text}',
                textfont={"size": 12},
                colorbar=dict(title="Продажи")
            ))
            
            fig_heatmap.update_layout(
                title="🗓️ Тепловая карта продаж",
                xaxis_title="День недели",
                yaxis_title="Неделя",
                height=500
            )
            
            st.plotly_chart(fig_heatmap, use_container_width=True, key="sales_heatmap")
            
            st.markdown("## 🏆 Топ-10 моделей по сегментам")
            
            segments_top_models = get_top_models_by_segment(df, selected_magazin)
            
            if segments_top_models:
                tabs = st.tabs([f"📦 {segment}" for segment in segments_top_models.keys()])
                
                for tab, (segment, top_models) in zip(tabs, segments_top_models.items()):
                    with tab:
                        if not top_models.empty:
                            display_df = top_models[['Model', 'Qty', 'Sum', 'Price']].rename(columns={
                                'Model': '🏷️ Модель',
                                'Qty': '📦 Количество',
                                'Sum': '💰 Выручка (ГРН)',
                                'Price': '💵 Средняя цена'
                            })
                            
                            st.dataframe(display_df, use_container_width=True, hide_index=True)
                        else:
                            st.info("🔍 Нет данных для этого сегмента")
            
            st.markdown("## 💡 Инсайты и рекомендации")
            
            insights, problems = generate_insights(df, forecast, selected_magazin, selected_segment)
            
            if problems:
                st.markdown("### 🚨 Выявленные проблемы:")
                for problem in problems:
                    st.markdown(f'<div class="problem-card">{problem}</div>', unsafe_allow_html=True)
            
            st.markdown("### 🎯 Рекомендации:")
            for insight in insights:
                st.markdown(f'<div class="insight-card">{insight}</div>', unsafe_allow_html=True)
            
            st.markdown("## 📋 Детальный прогноз по дням")
            
            forecast_display = forecast.tail(forecast_days).copy()
            segment_volatility = calculate_segment_volatility(df, selected_magazin, selected_segment)
            
            realistic, optimistic, pessimistic = get_forecast_scenarios(forecast_display, segment_volatility)
            
            detailed_forecast = pd.DataFrame({
                '📅 Дата': pd.to_datetime(forecast_display['ds']).dt.strftime('%Y-%m-%d (%A)'),
                '😰 Пессимистичный': pessimistic.round(0).astype(int),
                '🎯 Реальный': realistic.round(0).astype(int),
                '🚀 Оптимистичный': optimistic.round(0).astype(int),
                '📊 Тренд': forecast_display['trend'].round(0).astype(int)
            })
            
            st.dataframe(detailed_forecast, use_container_width=True, hide_index=True)
            
            st.markdown("## 📈 Дополнительная аналитика")
            
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                avg_daily_forecast = realistic.mean()
                st.metric(
                    "📊 Средние продажи/день",
                    f"{avg_daily_forecast:.0f}",
                    delta=f"{avg_daily_forecast - prophet_data['y'].tail(30).mean():.0f}"
                )
            
            with col2:
                total_forecast = realistic.sum()
                st.metric(
                    "📦 Общий прогноз",
                    f"{total_forecast:.0f}",
                    delta=f"{total_forecast - prophet_data['y'].tail(forecast_days).sum():.0f}"
                )
            
            with col3:
                # ИСПРАВЛЕНИЕ: безопасный расчет средней цены
                if len(filtered_df) > 0 and filtered_df['Qty'].sum() > 0:
                    avg_price = filtered_df['Sum'].sum() / filtered_df['Qty'].sum()
                else:
                    avg_price = 0
                
                forecast_revenue = total_forecast * avg_price
                st.metric(
                    "💰 Прогноз выручки",
                    f"{forecast_revenue:.0f} ГРН"
                )
            
            with col4:
                confidence_score = (1 - segment_volatility) * 100
                st.metric(
                    "🎯 Уверенность прогноза",
                    f"{confidence_score:.0f}%"
                )
            
            st.markdown("## 📥 Экспорт результатов")
            
            # === ИТОГОВЫЕ ВЫВОДЫ И РЕКОМЕНДАЦИИ ===
            st.markdown("---")
            st.markdown("## 🎓 Итоговые выводы и рекомендации")
            
            # Подготовка данных для анализа
            total_sales = filtered_df['Qty'].sum()
            total_revenue = filtered_df['Sum'].sum()
            avg_daily_sales = filtered_df.groupby('Datasales')['Qty'].sum().mean()
            
            # Тренд последних 30 дней
            last_30_days = filtered_df[filtered_df['Datasales'] >= filtered_df['Datasales'].max() - pd.Timedelta(days=30)]
            trend_last_month = last_30_days['Qty'].sum()
            
            # Прогноз
            forecast_total = realistic.sum()
            forecast_revenue = forecast_total * avg_price if avg_price > 0 else 0
            
            # Волатильность
            daily_volatility = prophet_data['y'].std() / prophet_data['y'].mean() if prophet_data['y'].mean() > 0 else 0
            
            # День недели анализ
            filtered_df_weekday = filtered_df.copy()
            filtered_df_weekday['Weekday'] = filtered_df_weekday['Datasales'].dt.dayofweek
            best_weekday = filtered_df_weekday.groupby('Weekday')['Qty'].sum().idxmax()
            weekday_names = ['Понедельник', 'Вторник', 'Среда', 'Четверг', 'Пятница', 'Суббота', 'Воскресенье']
            best_day_name = weekday_names[best_weekday]
            
            # Создаем три колонки для выводов
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.markdown("### 📊 Ключевые метрики")
                st.markdown(f"""
                <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 20px; border-radius: 10px; color: white;">
                    <p style="font-size: 14px; margin: 5px 0;"><strong>📦 Всего продано:</strong> {total_sales:.0f} шт</p>
                    <p style="font-size: 14px; margin: 5px 0;"><strong>💰 Общая выручка:</strong> {total_revenue:.0f} ГРН</p>
                    <p style="font-size: 14px; margin: 5px 0;"><strong>📈 Среднее/день:</strong> {avg_daily_sales:.1f} шт</p>
                    <p style="font-size: 14px; margin: 5px 0;"><strong>💵 Средняя цена:</strong> {avg_price:.2f} ГРН</p>
                    <p style="font-size: 14px; margin: 5px 0;"><strong>📅 Лучший день:</strong> {best_day_name}</p>
                </div>
                """, unsafe_allow_html=True)
            
            with col2:
                st.markdown("### 🔮 Прогнозные показатели")
                forecast_growth = ((forecast_total / prophet_data['y'].tail(forecast_days).sum()) - 1) * 100 if prophet_data['y'].tail(forecast_days).sum() > 0 else 0
                
                st.markdown(f"""
                <div style="background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%); padding: 20px; border-radius: 10px; color: white;">
                    <p style="font-size: 14px; margin: 5px 0;"><strong>📊 Прогноз ({forecast_days} дн):</strong> {forecast_total:.0f} шт</p>
                    <p style="font-size: 14px; margin: 5px 0;"><strong>💰 Прогноз выручки:</strong> {forecast_revenue:.0f} ГРН</p>
                    <p style="font-size: 14px; margin: 5px 0;"><strong>📈 Изменение:</strong> {forecast_growth:+.1f}%</p>
                    <p style="font-size: 14px; margin: 5px 0;"><strong>🎯 Уверенность:</strong> {confidence_score:.0f}%</p>
                    <p style="font-size: 14px; margin: 5px 0;"><strong>📊 Волатильность:</strong> {daily_volatility*100:.1f}%</p>
                </div>
                """, unsafe_allow_html=True)
            
            with col3:
                st.markdown("### ⚡ Статус модели")
                
                # Определяем качество модели
                if accuracy_metrics:
                    mape = accuracy_metrics['MAPE']
                    r2 = accuracy_metrics['R2']
                    
                    if mape < 10 and r2 > 0.8:
                        model_quality = "🟢 Отличное"
                        quality_color = "#2ecc71"
                    elif mape < 20 and r2 > 0.6:
                        model_quality = "🟡 Хорошее"
                        quality_color = "#f39c12"
                    else:
                        model_quality = "🔴 Удовлетворительное"
                        quality_color = "#e74c3c"
                    
                    st.markdown(f"""
                    <div style="background: {quality_color}; padding: 20px; border-radius: 10px; color: white;">
                        <p style="font-size: 16px; margin: 5px 0; font-weight: bold;">{model_quality}</p>
                        <p style="font-size: 14px; margin: 5px 0;"><strong>MAE:</strong> {accuracy_metrics['MAE']:.2f}</p>
                        <p style="font-size: 14px; margin: 5px 0;"><strong>MAPE:</strong> {mape:.1f}%</p>
                        <p style="font-size: 14px; margin: 5px 0;"><strong>R²:</strong> {r2:.3f}</p>
                    </div>
                    """, unsafe_allow_html=True)
                else:
                    st.info("Метрики недоступны")
            
            # Детальные выводы
            st.markdown("### 📝 Детальные выводы")
            
            conclusions = []
            
            # Анализ тренда
            if forecast_growth > 10:
                conclusions.append({
                    'emoji': '🚀',
                    'title': 'Позитивный тренд',
                    'text': f'Прогноз показывает рост продаж на {forecast_growth:.1f}%. Это отличная возможность для масштабирования.',
                    'type': 'success'
                })
            elif forecast_growth < -10:
                conclusions.append({
                    'emoji': '📉',
                    'title': 'Негативный тренд',
                    'text': f'Ожидается снижение продаж на {abs(forecast_growth):.1f}%. Требуются корректирующие действия.',
                    'type': 'error'
                })
            else:
                conclusions.append({
                    'emoji': '📊',
                    'title': 'Стабильный тренд',
                    'text': f'Продажи остаются стабильными с изменением {forecast_growth:+.1f}%.',
                    'type': 'info'
                })
            
            # Анализ волатильности
            if daily_volatility < 0.2:
                conclusions.append({
                    'emoji': '✅',
                    'title': 'Низкая волатильность',
                    'text': f'Волатильность {daily_volatility*100:.1f}% указывает на стабильный спрос. Планирование запасов упрощается.',
                    'type': 'success'
                })
            elif daily_volatility > 0.5:
                conclusions.append({
                    'emoji': '⚠️',
                    'title': 'Высокая волатильность',
                    'text': f'Волатильность {daily_volatility*100:.1f}% требует гибкого управления запасами и буферных резервов.',
                    'type': 'warning'
                })
            
            # Анализ точности модели
            if accuracy_metrics and accuracy_metrics['MAPE'] < 15:
                conclusions.append({
                    'emoji': '🎯',
                    'title': 'Высокая точность прогноза',
                    'text': f'MAPE {accuracy_metrics["MAPE"]:.1f}% говорит о высокой надежности прогноза. Можно уверенно использовать для планирования.',
                    'type': 'success'
                })
            elif accuracy_metrics and accuracy_metrics['MAPE'] > 25:
                conclusions.append({
                    'emoji': '⚡',
                    'title': 'Умеренная точность',
                    'text': f'MAPE {accuracy_metrics["MAPE"]:.1f}% указывает на необходимость дополнительного анализа внешних факторов.',
                    'type': 'warning'
                })
            
            # Анализ дня недели
            weekday_std = filtered_df_weekday.groupby('Weekday')['Qty'].sum().std()
            if weekday_std > avg_daily_sales * 0.3:
                conclusions.append({
                    'emoji': '📅',
                    'title': 'Неравномерность по дням недели',
                    'text': f'Значительная разница в продажах по дням ({best_day_name} - лучший). Оптимизируйте график работы персонала.',
                    'type': 'info'
                })
            
            # Отображение выводов
            for conclusion in conclusions:
                if conclusion['type'] == 'success':
                    st.success(f"{conclusion['emoji']} **{conclusion['title']}**: {conclusion['text']}")
            
            
            # Стратегические рекомендации
            st.markdown("### 🎯 Стратегические рекомендации")
            
            recommendations_strategic = []
            
            # Рекомендации по закупкам
            optimal_stock = forecast_total * 1.2  # +20% буфер
            recommendations_strategic.append({
                'category': '📦 Управление запасами',
                'items': [
                    f"Оптимальный запас на {forecast_days} дней: {optimal_stock:.0f} единиц",
                    f"Средний дневной запас: {optimal_stock/forecast_days:.0f} единиц",
                    "Организуйте систему «точно вовремя» для быстроходных товаров" if daily_volatility < 0.3 else "Держите буферный запас 30-40% из-за высокой волатильности"
                ]
            })
            
            # Рекомендации по ценообразованию
            price_recommendations = []
            if len(filtered_df) >= 30:
                recent_price = filtered_df.tail(30)['Sum'].sum() / filtered_df.tail(30)['Qty'].sum() if filtered_df.tail(30)['Qty'].sum() > 0 else 0
                older_price = filtered_df.head(30)['Sum'].sum() / filtered_df.head(30)['Qty'].sum() if filtered_df.head(30)['Qty'].sum() > 0 else 0
                
                if recent_price > older_price * 1.1:
                    price_recommendations.append("Цены растут - следите за реакцией спроса")
                elif recent_price < older_price * 0.9:
                    price_recommendations.append("Цены падают - возможна ценовая конкуренция")
                else:
                    price_recommendations.append("Цены стабильны - хорошо для планирования")
            
            price_recommendations.append(f"Текущая средняя цена: {avg_price:.2f} ГРН")
            price_recommendations.append("Тестируйте эластичность спроса с A/B тестированием цен")
            
            recommendations_strategic.append({
                'category': '💰 Ценообразование',
                'items': price_recommendations
            })
            
            # Рекомендации по маркетингу
            marketing_recommendations = []
            
            if best_weekday in [5, 6]:  # Суббота, Воскресенье
                marketing_recommendations.append("Пиковые продажи в выходные - усильте маркетинг в четверг-пятницу")
            else:
                marketing_recommendations.append(f"Пик продаж в {best_day_name} - планируйте акции на этот день")
            
            weak_days = filtered_df_weekday.groupby('Weekday')['Qty'].sum()
            if weak_days.min() < weak_days.mean() * 0.7:
                marketing_recommendations.append("Проводите акции «счастливые часы» в слабые дни недели")
            
            if forecast_growth < 0:
                marketing_recommendations.append("Запустите промо-кампанию для стимулирования спроса")
            
            recommendations_strategic.append({
                'category': '📢 Маркетинг и продажи',
                'items': marketing_recommendations
            })
            
            # Рекомендации по персоналу
            staff_recommendations = [
                f"Максимум персонала в {best_day_name}",
                "Гибкий график для оптимизации затрат на ФОТ",
                f"Планируйте {int(optimal_stock/(forecast_days*8)):.0f} транзакций в час в пиковые дни"
            ]
            
            recommendations_strategic.append({
                'category': '👥 Управление персоналом',
                'items': staff_recommendations
            })
            
            # Отображение стратегических рекомендаций
            for rec in recommendations_strategic:
                with st.expander(f"### {rec['category']}", expanded=True):
                    for item in rec['items']:
                        st.markdown(f"- {item}")
            
            # Финальная рекомендация
            st.markdown("### 🎖️ Приоритетное действие")
            
            if forecast_growth > 15:
                priority = "🚀 **Масштабирование**: Увеличьте закупки на 20-30% и подготовьте логистику к росту спроса"
                priority_color = "#2ecc71"
            elif forecast_growth < -15:
                priority = "🔥 **Антикризисные меры**: Срочно запустите маркетинговую кампанию и пересмотрите ассортимент"
                priority_color = "#e74c3c"
            elif daily_volatility > 0.5:
                priority = "⚖️ **Стабилизация**: Сосредоточьтесь на сглаживании колебаний спроса через промо-активности"
                priority_color = "#f39c12"
            else:
                priority = "📊 **Оптимизация**: Продолжайте текущую стратегию с фокусом на увеличение маржинальности"
                priority_color = "#3498db"
            
            st.markdown(f"""
            <div style="background: {priority_color}; padding: 20px; border-radius: 10px; color: white; font-size: 18px; text-align: center; font-weight: bold; margin: 20px 0;">
                {priority}
            </div>
            """, unsafe_allow_html=True)
            
            st.markdown("---")
            st.markdown("## 📥 Экспорт результатов")
            
            # === МАРКЕТИНГОВАЯ АНАЛИТИКА ===
            st.markdown("---")
            st.markdown("## 📢 Маркетинговая аналитика")
            
            st.markdown("### 🎯 Ключевые маркетинговые метрики")
            
            # Расчет маркетинговых метрик
            # 1. Customer Lifetime Value (приблизительный)
            avg_transaction = filtered_df['Sum'].sum() / len(filtered_df) if len(filtered_df) > 0 else 0
            transactions_per_day = len(filtered_df) / ((filtered_df['Datasales'].max() - filtered_df['Datasales'].min()).days + 1)
            
            # 2. ABC анализ товаров
            product_analysis = filtered_df.groupby('Model').agg({
                'Qty': 'sum',
                'Sum': 'sum'
            }).reset_index()
            product_analysis = product_analysis.sort_values('Sum', ascending=False)
            product_analysis['Cumulative_Revenue'] = product_analysis['Sum'].cumsum()
            product_analysis['Cumulative_Percent'] = (product_analysis['Cumulative_Revenue'] / product_analysis['Sum'].sum()) * 100
            
            # Классификация ABC
            product_analysis['Category'] = 'C'
            product_analysis.loc[product_analysis['Cumulative_Percent'] <= 80, 'Category'] = 'A'
            product_analysis.loc[(product_analysis['Cumulative_Percent'] > 80) & (product_analysis['Cumulative_Percent'] <= 95), 'Category'] = 'B'
            
            # 3. Анализ жизненного цикла товара
            first_sale = filtered_df.groupby('Model')['Datasales'].min()
            last_sale = filtered_df.groupby('Model')['Datasales'].max()
            product_lifecycle = pd.DataFrame({
                'First_Sale': first_sale,
                'Last_Sale': last_sale,
                'Days_Active': (last_sale - first_sale).dt.days
            })
            
            # 4. Conversion rate (условный - продажи vs просмотры)
            daily_products = filtered_df.groupby('Datasales')['Art'].nunique().mean()
            daily_sales = filtered_df.groupby('Datasales')['Qty'].sum().mean()
            conversion_rate = (daily_sales / daily_products) if daily_products > 0 else 0
            
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                st.metric(
                    "💳 Средний чек",
                    f"{avg_transaction:.0f} ГРН",
                    help="Средняя сумма одной транзакции"
                )
            
            with col2:
                st.metric(
                    "🔄 Транзакций/день",
                    f"{transactions_per_day:.1f}",
                    help="Среднее количество транзакций в день"
                )
            
            with col3:
                st.metric(
                    "📊 Conversion Rate",
                    f"{conversion_rate:.1f}x",
                    help="Соотношение продаж к уникальным товарам"
                )
            
            with col4:
                category_a_count = (product_analysis['Category'] == 'A').sum()
                st.metric(
                    "⭐ Товары категории A",
                    f"{category_a_count}",
                    help="Товары, дающие 80% выручки"
                )
            
            # ABC Анализ - таблица по категориям
            st.markdown("### 📊 ABC-анализ товаров (Правило Парето)")
            
            # Создаем сводную таблицу по категориям
            abc_summary = product_analysis.groupby('Category').agg({
                'Model': 'count',
                'Sum': 'sum'
            }).reset_index()
            abc_summary.columns = ['Категория', 'Количество позиций', 'Доход (ГРН)']
            
            # Добавляем процент от общего дохода
            total_revenue = abc_summary['Доход (ГРН)'].sum()
            abc_summary['Доля от общего дохода %'] = (abc_summary['Доход (ГРН)'] / total_revenue * 100).round(2)
            
            # Сортируем по категориям A, B, C
            category_order = {'A': 0, 'B': 1, 'C': 2}
            abc_summary['sort_key'] = abc_summary['Категория'].map(category_order)
            abc_summary = abc_summary.sort_values('sort_key').drop('sort_key', axis=1)
            
            # Форматирование таблицы с эмодзи
            abc_summary['Категория'] = abc_summary['Категория'].map({
                'A': '⭐ Категория A (80% выручки)',
                'B': '🔶 Категория B (15% выручки)',
                'C': '🔻 Категория C (5% выручки)'
            })
            
            st.dataframe(
                abc_summary.style.format({
                    'Количество позиций': '{:.0f}',
                    'Доход (ГРН)': '{:.0f}',
                    'Доля от общего дохода %': '{:.2f}%'
                }).background_gradient(subset=['Доход (ГРН)'], cmap='Greens'),
                use_container_width=True,
                hide_index=True
            )
            
            # Топ товары категории A
            st.markdown("#### ⭐ Топ-10 товаров категории A (приносят наибольшую выручку)")
            
            top_a_products = product_analysis[product_analysis['Category'] == 'A'].head(10)
            top_a_display = top_a_products[['Model', 'Qty', 'Sum']].copy()
            top_a_display['Revenue_Share_%'] = (top_a_display['Sum'] / filtered_df['Sum'].sum() * 100).round(2)
            top_a_display = top_a_display.rename(columns={
                'Model': '🏷️ Модель',
                'Qty': '📦 Продано',
                'Sum': '💰 Выручка (ГРН)',
                'Revenue_Share_%': '📊 Доля выручки %'
            })
            
            st.dataframe(
                top_a_display.style.format({
                    '📦 Продано': '{:.0f}',
                    '💰 Выручка (ГРН)': '{:.0f}',
                    '📊 Доля выручки %': '{:.2f}%'
                }).background_gradient(subset=['💰 Выручка (ГРН)'], cmap='Greens'),
                use_container_width=True,
                hide_index=True
            )
            
            # ТОП-20 товаров по категориям
            st.markdown("### 🏆 ТОП-20 товаров по категориям")
            
            # Подготовка данных для анализа тренда
            # Разделяем данные на 2 периода для сравнения
            mid_date = filtered_df['Datasales'].min() + (filtered_df['Datasales'].max() - filtered_df['Datasales'].min()) / 2
            
            period1 = filtered_df[filtered_df['Datasales'] < mid_date]
            period2 = filtered_df[filtered_df['Datasales'] >= mid_date]
            
            # Продажи по периодам
            sales_period1 = period1.groupby('Model')['Qty'].sum()
            sales_period2 = period2.groupby('Model')['Qty'].sum()
            
            # Общие продажи
            total_sales_by_model = filtered_df.groupby('Model').agg({
                'Qty': 'sum',
                'Sum': 'sum'
            })
            
            # Расчет изменения
            trend_data = pd.DataFrame({
                'Model': total_sales_by_model.index,
                'Total_Qty': total_sales_by_model['Qty'].values,
                'Total_Revenue': total_sales_by_model['Sum'].values,
                'Period1_Qty': sales_period1.reindex(total_sales_by_model.index, fill_value=0).values,
                'Period2_Qty': sales_period2.reindex(total_sales_by_model.index, fill_value=0).values
            })
            
            # Расчет процента изменения
            trend_data["Change_%"] = np.where(
                trend_data["Period1_Qty"] > 0,
                ((trend_data["Period2_Qty"] - trend_data["Period1_Qty"]) / trend_data["Period1_Qty"]) * 100,
                0
            )
            # Расчет процента изменения
            trend_data["Change_%"] = np.where(
                trend_data["Period1_Qty"] > 0,
                ((trend_data["Period2_Qty"] - trend_data["Period1_Qty"]) / trend_data["Period1_Qty"]) * 100,
                0
            )
            # Расчет процента изменения
            trend_data["Change_%"] = np.where(
                trend_data["Period1_Qty"] > 0,
                ((trend_data["Period2_Qty"] - trend_data["Period1_Qty"]) / trend_data["Period1_Qty"]) * 100,
                0
            )
            
            # Стабильность (чем меньше изменение, тем стабильнее)
            trend_data['Stability_Score'] = 100 - abs(trend_data['Change_%']).clip(upper=100)
            
            # Создаем вкладки
            tab1, tab2, tab3 = st.tabs(["🏆 ТОП-20 Лучших", "📊 ТОП-20 Стабильных", "📉 ТОП-20 Падение"])
            
            # ТОП-20 ЛУЧШИХ
            with tab1:
                st.markdown("#### 🏆 ТОП-20 моделей по выручке")
                
                top_20_best = trend_data.nlargest(20, 'Total_Revenue').copy()
                top_20_best['Avg_Price'] = top_20_best['Total_Revenue'] / top_20_best['Total_Qty']
                
                # График
                fig_best = go.Figure()
                
                fig_best.add_trace(go.Bar(
                    y=top_20_best['Model'],
                    x=top_20_best['Total_Revenue'],
                    orientation='h',
                    marker=dict(
                        color=top_20_best['Total_Revenue'],
                        colorscale='Greens',
                        showscale=True,
                        colorbar=dict(title="Выручка<br>ГРН")
                    ),
                    text=top_20_best['Total_Revenue'].apply(lambda x: f'{x:.0f}'),
                    textposition='outside',
                    hovertemplate='<b>%{y}</b><br>Выручка: %{x:.0f} ГРН<br><extra></extra>'
                ))
                
                fig_best.update_layout(
                    title="ТОП-20 моделей по выручке",
                    xaxis_title="Выручка (ГРН)",
                    yaxis_title="Модель",
                    height=600,
                    yaxis={'categoryorder': 'total ascending'}
                )
                
                st.plotly_chart(fig_best, use_container_width=True, key="top20_best")
                
                # Таблица
                st.markdown("##### 📋 Детальная информация")
                
                display_best = top_20_best[['Model', 'Total_Qty', 'Total_Revenue', 'Avg_Price', 'Change_%']].copy()
                display_best = display_best.reset_index(drop=True)
                display_best.index = display_best.index + 1
                display_best = display_best.rename(columns={
                    'Model': '🏷️ Модель',
                    'Total_Qty': '📦 Продано',
                    'Total_Revenue': '💰 Выручка (ГРН)',
                    'Avg_Price': '💵 Средняя цена',
                    'Change_%': '📈 Изменение %'
                })
                
                st.dataframe(
                    display_best.style.format({
                        '📦 Продано': '{:.0f}',
                        '💰 Выручка (ГРН)': '{:.0f}',
                        '💵 Средняя цена': '{:.2f}',
                        '📈 Изменение %': '{:+.1f}%'
                    }).background_gradient(subset=['💰 Выручка (ГРН)'], cmap='Greens')
                    .background_gradient(subset=['📈 Изменение %'], cmap='RdYlGn', vmin=-50, vmax=50),
                    use_container_width=True
                )
                
                # Инсайты
                st.success(f"💎 **Лидер продаж**: {top_20_best.iloc[0]['Model']} - выручка {top_20_best.iloc[0]['Total_Revenue']:.0f} ГРН")
                
                growth_products = top_20_best[top_20_best['Change_%'] > 20]
                if len(growth_products) > 0:
                    st.info(f"🚀 **Растущие хиты**: {len(growth_products)} товаров показывают рост более 20%")
            
            # ТОП-20 СТАБИЛЬНЫХ
            with tab2:
                st.markdown("#### 📊 ТОП-20 самых стабильных моделей")
                st.caption("Товары с минимальными колебаниями продаж между периодами")
                
                # Фильтруем только те товары, которые продавались в обоих периодах
                stable_products = trend_data[(trend_data['Period1_Qty'] > 0) & (trend_data['Period2_Qty'] > 0)].copy()
                top_20_stable = stable_products.nlargest(20, 'Stability_Score')
                top_20_stable['Avg_Price'] = top_20_stable['Total_Revenue'] / top_20_stable['Total_Qty']
                
                # График
                fig_stable = go.Figure()
                
                fig_stable.add_trace(go.Bar(
                    y=top_20_stable['Model'],
                    x=top_20_stable['Stability_Score'],
                    orientation='h',
                    marker=dict(
                        color=top_20_stable['Stability_Score'],
                        colorscale='Blues',
                        showscale=True,
                        colorbar=dict(title="Стабильность<br>%")
                    ),
                    text=top_20_stable['Stability_Score'].apply(lambda x: f'{x:.1f}'),
                    textposition='outside',
                    hovertemplate='<b>%{y}</b><br>Стабильность: %{x:.1f}%<br><extra></extra>'
                ))
                
                fig_stable.update_layout(
                    title="ТОП-20 самых стабильных моделей",
                    xaxis_title="Индекс стабильности (%)",
                    yaxis_title="Модель",
                    height=600,
                    yaxis={'categoryorder': 'total ascending'}
                )
                
                st.plotly_chart(fig_stable, use_container_width=True, key="top20_stable")
                
                # Таблица
                st.markdown("##### 📋 Детальная информация")
                
                display_stable = top_20_stable[['Model', 'Total_Qty', 'Total_Revenue', 'Avg_Price', 'Change_%', 'Stability_Score']].copy()
                display_stable = display_stable.reset_index(drop=True)
                display_stable.index = display_stable.index + 1
                display_stable = display_stable.rename(columns={
                    'Model': '🏷️ Модель',
                    'Total_Qty': '📦 Продано',
                    'Total_Revenue': '💰 Выручка (ГРН)',
                    'Avg_Price': '💵 Средняя цена',
                    'Change_%': '📈 Изменение %',
                    'Stability_Score': '📊 Стабильность %'
                })
                
                st.dataframe(
                    display_stable.style.format({
                        '📦 Продано': '{:.0f}',
                        '💰 Выручка (ГРН)': '{:.0f}',
                        '💵 Средняя цена': '{:.2f}',
                        '📈 Изменение %': '{:+.1f}%',
                        '📊 Стабильность %': '{:.1f}%'
                    }).background_gradient(subset=['📊 Стабильность %'], cmap='Blues'),
                    use_container_width=True
                )
                
                # Инсайты
                st.success(f"🎯 **Самый стабильный**: {top_20_stable.iloc[0]['Model']} - стабильность {top_20_stable.iloc[0]['Stability_Score']:.1f}%")
                st.info(f"💡 **Рекомендация**: Стабильные товары идеальны для постоянного наличия на складе")
            
            # ТОП-20 ПАДЕНИЕ
            with tab3:
                st.markdown("#### 📉 ТОП-20 моделей с наибольшим падением")
                st.caption("Товары, показывающие снижение продаж")
                
                # Только товары с падением
                declining_products = trend_data[trend_data['Change_%'] < 0].copy()
                
                if len(declining_products) > 0:
                    top_20_declining = declining_products.nsmallest(20, 'Change_%')
                    top_20_declining['Avg_Price'] = top_20_declining['Total_Revenue'] / top_20_declining['Total_Qty']
                    
                    # График
                    fig_decline = go.Figure()
                    
                    fig_decline.add_trace(go.Bar(
                        y=top_20_declining['Model'],
                        x=top_20_declining['Change_%'],
                        orientation='h',
                        marker=dict(
                            color=top_20_declining['Change_%'],
                            colorscale='Reds_r',
                            showscale=True,
                            colorbar=dict(title="Падение<br>%")
                        ),
                        text=top_20_declining['Change_%'].apply(lambda x: f'{x:.1f}%'),
                        textposition='outside',
                        hovertemplate='<b>%{y}</b><br>Падение: %{x:.1f}%<br><extra></extra>'
                    ))
                    
                    fig_decline.update_layout(
                        title="ТОП-20 моделей с наибольшим падением продаж",
                        xaxis_title="Изменение (%)",
                        yaxis_title="Модель",
                        height=600,
                        yaxis={'categoryorder': 'total descending'}
                    )
                    
                    st.plotly_chart(fig_decline, use_container_width=True, key="top20_decline")
                    
                    # Таблица
                    st.markdown("##### 📋 Детальная информация")
                    
                    display_decline = top_20_declining[['Model', 'Total_Qty', 'Total_Revenue', 'Period1_Qty', 'Period2_Qty', 'Change_%']].copy()
                    display_decline = display_decline.reset_index(drop=True)
                    display_decline.index = display_decline.index + 1
                    display_decline = display_decline.rename(columns={
                        'Model': '🏷️ Модель',
                        'Total_Qty': '📦 Всего продано',
                        'Total_Revenue': '💰 Выручка (ГРН)',
                        'Period1_Qty': '📊 1-й период',
                        'Period2_Qty': '📊 2-й период',
                        'Change_%': '📉 Падение %'
                    })
                    
                    st.dataframe(
                        display_decline.style.format({
                            '📦 Всего продано': '{:.0f}',
                            '💰 Выручка (ГРН)': '{:.0f}',
                            '📊 1-й период': '{:.0f}',
                            '📊 2-й период': '{:.0f}',
                            '📉 Падение %': '{:.1f}%'
                        }).background_gradient(subset=['📉 Падение %'], cmap='Reds_r'),
                        use_container_width=True
                    )
                    
                    # Алерты и рекомендации
                    critical_decline = top_20_declining[top_20_declining['Change_%'] < -50]
                    
                    if len(critical_decline) > 0:
                        st.error(f"🚨 **КРИТИЧНО**: {len(critical_decline)} товаров с падением более 50%!")
                    
                    st.warning(f"⚠️ **Проблемный товар**: {top_20_declining.iloc[0]['Model']} - падение {top_20_declining.iloc[0]['Change_%']:.1f}%")
                    
                    st.markdown("##### 💡 Рекомендации по проблемным товарам:")
                    recommendations = [
                        "🎯 Провести анализ причин падения (конкуренция, цена, актуальность)",
                        "🔥 Запустить промо-акции со скидками 20-30%",
                        "📢 Усилить маркетинг и рекламу",
                        "💡 Рассмотреть обновление ассортимента или замену товара",
                        "📦 Снизить закупки до стабилизации ситуации"
                    ]
                    
                    for rec in recommendations:
                        st.markdown(f'<div class="insight-card">{rec}</div>', unsafe_allow_html=True)
                    
                else:
                    st.success("✅ Отличные новости! Нет товаров с падением продаж")
                    st.balloons()
            
            # Маркетинговые рекомендации
            st.markdown("### 🎯 Маркетинговые рекомендации")
            
            # Табличные метрики
            st.markdown("#### 📊 Ключевые метрики маркетинга")
            
            # Расчет метрик
            top_5_revenue = trend_data.nlargest(5, 'Total_Revenue')
            total_revenue_all = trend_data['Total_Revenue'].sum()
            top_5_share = (top_5_revenue['Total_Revenue'].sum() / total_revenue_all) * 100
            top_10_share = (trend_data.nlargest(10, 'Total_Revenue')['Total_Revenue'].sum() / total_revenue_all) * 100
            
            declining_count = len(trend_data[trend_data['Change_%'] < -20])
            growing_count = len(trend_data[trend_data['Change_%'] > 20])
            stable_count = len(trend_data[(trend_data['Change_%'] >= -20) & (trend_data['Change_%'] <= 20)])
            
            avg_transaction = filtered_df['Sum'].sum() / len(filtered_df) if len(filtered_df) > 0 else 0
            avg_price = filtered_df['Price'].mean()
            avg_qty_per_transaction = filtered_df.groupby('Datasales')['Qty'].sum().mean()
            
            # Создание таблицы метрик
            metrics_data = {
                'Метрика': [
                    'Доля ТОП-5 товаров',
                    'Доля ТОП-10 товаров',
                    'Товаров с ростом >20%',
                    'Товаров с падением >20%',
                    'Стабильных товаров',
                    'Средний чек',
                    'Средняя цена товара',
                    'Средние продажи/день'
                ],
                'Значение': [
                    f'{top_5_share:.1f}%',
                    f'{top_10_share:.1f}%',
                    f'{growing_count} шт.',
                    f'{declining_count} шт.',
                    f'{stable_count} шт.',
                    f'{avg_transaction:.0f} ГРН',
                    f'{avg_price:.0f} ГРН',
                    f'{avg_qty_per_transaction:.1f} шт.'
                ],
                'Статус': [
                    '⚠️ Риск' if top_5_share > 50 else '✅ Норма',
                    '⚠️ Риск' if top_10_share > 70 else '✅ Норма',
                    '✅ Рост' if growing_count > 10 else '⚠️ Мало',
                    '🚨 Критично' if declining_count > 10 else ('⚠️ Внимание' if declining_count > 0 else '✅ OK'),
                    '✅ Хорошо' if stable_count > 20 else '⚠️ Мало',
                    '⚠️ Низкий' if avg_transaction < filtered_df['Sum'].mean() * 0.8 else '✅ Норма',
                    '✅ Норма',
                    '✅ Норма'
                ]
            }
            
            metrics_df = pd.DataFrame(metrics_data)
            st.dataframe(metrics_df, use_container_width=True, hide_index=True)
            
            # Выводы и ивенты
            st.markdown("#### 💡 Ключевые выводы")
            
            conclusions = []
            events = []
            
            # Анализ концентрации
            if top_5_share > 50:
                conclusions.append("🎯 **Высокая концентрация выручки** - зависимость от узкого ассортимента")
                events.append(f"📌 СОБЫТИЕ: ТОП-5 товаров генерируют {top_5_share:.1f}% выручки")
            
            # Анализ динамики
            if declining_count > 10:
                conclusions.append("📉 **Критическая динамика** - массовое падение продаж товаров")
                events.append(f"🚨 СОБЫТИЕ: {declining_count} товаров теряют продажи (>20% падения)")
            elif declining_count > 0:
                conclusions.append("⚠️ **Проблемные позиции** - есть товары с падением продаж")
                events.append(f"⚠️ СОБЫТИЕ: {declining_count} товаров требуют корректировки стратегии")
            
            if growing_count > 10:
                conclusions.append("📈 **Позитивный тренд** - значительное количество растущих товаров")
                events.append(f"✅ СОБЫТИЕ: {growing_count} товаров демонстрируют рост (>20%)")
            
            # Анализ среднего чека
            if avg_transaction < filtered_df['Sum'].mean() * 0.8:
                conclusions.append("💳 **Низкий средний чек** - потенциал увеличения через cross-sell")
                events.append(f"💰 СОБЫТИЕ: Средний чек {avg_transaction:.0f} ГРН ниже оптимального")
            
            # Дополнительные события
            if total_revenue_all > 0:
                top_product = trend_data.nlargest(1, 'Total_Revenue').iloc[0]
                events.append(f"🏆 ЛИДЕР: {top_product['Model']} - {top_product['Total_Revenue']:.0f} ГРН выручки")
            
            # Отображение выводов
            for conclusion in conclusions:
                st.markdown(f'<div class="insight-card">{conclusion}</div>', unsafe_allow_html=True)
            
            st.markdown("#### 📅 Важные события периода")
            for event in events:
                st.info(event)
            
            # Инсайты
            marketing_insights = []
            
            if top_5_share > 50:
                marketing_insights.append({
                    'type': 'warning',
                    'title': '⚠️ Высокая концентрация',
                    'text': f"ТОП-5 товаров дают {top_5_share:.1f}% выручки. Это риск! Диверсифицируйте портфель."
                })
            else:
                marketing_insights.append({
                    'type': 'success',
                    'title': '✅ Сбалансированный портфель',
                    'text': f"ТОП-5 товаров дают {top_5_share:.1f}% выручки. Хорошая диверсификация."
                })
            
            if declining_count > 10:
                marketing_insights.append({
                    'type': 'error',
                    'title': '🚨 Критическая ситуация',
                    'text': f"{declining_count} товаров с падением >20%. Срочно пересмотреть стратегию!"
                })
            elif declining_count > 0:
                marketing_insights.append({
                    'type': 'warning',
                    'title': '📉 Требуется внимание',
                    'text': f"{declining_count} товаров с падением >20%. Проведите анализ и акции."
                })
            
            if avg_transaction < filtered_df['Sum'].mean() * 0.8:
                marketing_insights.append({
                    'type': 'info',
                    'title': '💳 Низкий средний чек',
                    'text': f"Средний чек {avg_transaction:.0f} ГРН. Внедрите cross-sell и бандлы товаров."
                })
            
            st.markdown("#### 🎯 Рекомендации")
            for insight in marketing_insights:
                if insight['type'] == 'success':
                    st.success(f"**{insight['title']}**: {insight['text']}")
                elif insight['type'] == 'warning':
                    st.warning(f"**{insight['title']}**: {insight['text']}")
                elif insight['type'] == 'error':
                    st.error(f"**{insight['title']}**: {insight['text']}")
                else:
                    st.info(f"**{insight['title']}**: {insight['text']}")
            
            # Конкретные действия
            st.markdown("#### 📋 План действий на ближайшие 30 дней")
            
            top_20_best_count = len(trend_data.nlargest(20, 'Total_Revenue'))
            declining_count_action = len(trend_data[trend_data['Change_%'] < -20])
            
            action_plan = [
                f"1️⃣ **ТОП товары**: Увеличить бюджет на рекламу ТОП-{top_20_best_count} товаров на 30%",
                f"2️⃣ **Падающие товары**: Провести распродажу товаров с падением со скидкой 20-30%",
                f"3️⃣ **Cross-sell**: Создать 5 товарных бандлов для увеличения среднего чека",
                f"4️⃣ **Стабильные товары**: Обеспечить постоянное наличие стабильных позиций на складе",
                f"5️⃣ **Мониторинг**: Еженедельно отслеживать динамику падающих товаров"
            ]
            
            for action in action_plan:
                st.markdown(f'<div class="insight-card">{action}</div>', unsafe_allow_html=True)
            
            # Анализ эластичности спроса
            st.markdown("### 📐 Анализ эластичности спроса")
            st.markdown("Оценка чувствительности спроса к изменению цены (анализ по всему датасету)")
            
            # Расчет эластичности для товаров с достаточными данными
            elasticity_data = []
            
            # Используем весь датасет df вместо filtered_df
            all_models = df['Model'].unique()
            
            # Анализируем все модели с достаточным количеством данных
            for model in all_models:
                model_data = df[df['Model'] == model].copy()
                
                if len(model_data) >= 10:  # Минимум 10 записей для анализа
                    # Группировка по ценовым диапазонам
                    try:
                        model_data['Price_Group'] = pd.qcut(model_data['Price'], q=3, labels=['Низкая', 'Средняя', 'Высокая'], duplicates='drop')
                    except:
                        # Если не получается разбить на 3 группы, пробуем на 2
                        try:
                            model_data['Price_Group'] = pd.qcut(model_data['Price'], q=2, labels=['Низкая', 'Высокая'], duplicates='drop')
                        except:
                            continue
                    
                    price_analysis = model_data.groupby('Price_Group').agg({
                        'Price': 'mean',
                        'Qty': 'sum'
                    }).reset_index()
                    
                    if len(price_analysis) >= 2:
                        # Простой расчет эластичности между крайними группами
                        if price_analysis.iloc[0]['Price'] != price_analysis.iloc[-1]['Price']:
                            price_change_pct = ((price_analysis.iloc[-1]['Price'] - price_analysis.iloc[0]['Price']) / 
                                              price_analysis.iloc[0]['Price']) * 100
                            qty_change_pct = ((price_analysis.iloc[-1]['Qty'] - price_analysis.iloc[0]['Qty']) / 
                                            price_analysis.iloc[0]['Qty']) * 100
                            
                            if price_change_pct != 0:
                                elasticity = qty_change_pct / price_change_pct
                                
                                # Классификация эластичности
                                if abs(elasticity) > 1:
                                    elasticity_type = "Эластичный"
                                    recommendation = "Снижение цены увеличит выручку"
                                elif abs(elasticity) < 1:
                                    elasticity_type = "Неэластичный"
                                    recommendation = "Повышение цены увеличит выручку"
                                else:
                                    elasticity_type = "Единичный"
                                    recommendation = "Цена оптимальна"
                                
                                total_revenue = model_data['Sum'].sum()
                                avg_price = model_data['Price'].mean()
                                total_qty = model_data['Qty'].sum()
                                
                                elasticity_data.append({
                                    'Model': model,
                                    'Elasticity': elasticity,
                                    'Type': elasticity_type,
                                    'Avg_Price': avg_price,
                                    'Total_Revenue': total_revenue,
                                    'Total_Qty': total_qty,
                                    'Price_Change_%': price_change_pct,
                                    'Qty_Change_%': qty_change_pct,
                                    'Recommendation': recommendation
                                })
            
            if len(elasticity_data) > 0:
                elasticity_df = pd.DataFrame(elasticity_data)
                elasticity_df = elasticity_df.sort_values('Total_Revenue', ascending=False)
                
                # Метрики эластичности
                col1, col2, col3, col4 = st.columns(4)
                
                elastic_count = len(elasticity_df[elasticity_df['Type'] == 'Эластичный'])
                inelastic_count = len(elasticity_df[elasticity_df['Type'] == 'Неэластичный'])
                unit_count = len(elasticity_df[elasticity_df['Type'] == 'Единичный'])
                
                with col1:
                    st.metric("📊 Проанализировано товаров", len(elasticity_df))
                with col2:
                    st.metric("⚡ Эластичных", elastic_count)
                with col3:
                    st.metric("🔒 Неэластичных", inelastic_count)
                with col4:
                    st.metric("⚖️ Единичных", unit_count)
                
                # График эластичности
                st.markdown("#### 📈 Распределение коэффициентов эластичности")
                
                fig_elasticity = go.Figure()
                
                colors = elasticity_df['Type'].map({
                    'Эластичный': '#ff6b6b',
                    'Неэластичный': '#51cf66',
                    'Единичный': '#ffd43b'
                })
                
                fig_elasticity.add_trace(go.Bar(
                    y=elasticity_df['Model'].head(20),
                    x=elasticity_df['Elasticity'].head(20),
                    orientation='h',
                    marker=dict(
                        color=colors.head(20),
                        line=dict(color='white', width=1)
                    ),
                    text=elasticity_df['Elasticity'].head(20).apply(lambda x: f'{x:.2f}'),
                    textposition='outside',
                    hovertemplate='<b>%{y}</b><br>Эластичность: %{x:.2f}<extra></extra>'
                ))
                
                fig_elasticity.add_vline(x=-1, line_dash="dash", line_color="red", 
                                        annotation_text="Граница эластичности")
                fig_elasticity.add_vline(x=1, line_dash="dash", line_color="red")
                
                fig_elasticity.update_layout(
                    title="ТОП-20 товаров по коэффициенту эластичности (весь датасет)",
                    xaxis_title="Коэффициент эластичности",
                    yaxis_title="Модель",
                    height=600,
                    showlegend=False
                )
                
                st.plotly_chart(fig_elasticity, use_container_width=True)
                
                # Таблица с рекомендациями
                st.markdown("#### 📋 Детальный анализ и рекомендации")
                
                display_elasticity = elasticity_df.head(20)[['Model', 'Type', 'Elasticity', 
                                                             'Avg_Price', 'Total_Revenue', 'Total_Qty',
                                                             'Price_Change_%', 'Qty_Change_%', 
                                                             'Recommendation']].copy()
                
                display_elasticity = display_elasticity.rename(columns={
                    'Model': '🏷️ Модель',
                    'Type': '📊 Тип',
                    'Elasticity': '📐 Эластичность',
                    'Avg_Price': '💰 Средняя цена',
                    'Total_Revenue': '💵 Выручка',
                    'Total_Qty': '📦 Продано шт.',
                    'Price_Change_%': '📈 Изм. цены %',
                    'Qty_Change_%': '📊 Изм. объема %',
                    'Recommendation': '💡 Рекомендация'
                })
                
                st.dataframe(
                    display_elasticity.style.format({
                        '📐 Эластичность': '{:.2f}',
                        '💰 Средняя цена': '{:.0f} ГРН',
                        '💵 Выручка': '{:.0f} ГРН',
                        '📦 Продано шт.': '{:.0f}',
                        '📈 Изм. цены %': '{:.1f}%',
                        '📊 Изм. объема %': '{:.1f}%'
                    }).applymap(
                        lambda x: 'background-color: #ffebee' if x == 'Эластичный' else 
                                 ('background-color: #e8f5e9' if x == 'Неэластичный' else 
                                  ('background-color: #fff9c4' if x == 'Единичный' else '')),
                        subset=['📊 Тип']
                    ),
                    use_container_width=True
                )
                
                # Стратегические рекомендации
                st.markdown("#### 🎯 Стратегические рекомендации по ценообразованию")
                
                elastic_revenue = elasticity_df[elasticity_df['Type'] == 'Эластичный']['Total_Revenue'].sum()
                inelastic_revenue = elasticity_df[elasticity_df['Type'] == 'Неэластичный']['Total_Revenue'].sum()
                total_analyzed_revenue = elastic_revenue + inelastic_revenue
                
                pricing_recommendations = []
                
                if elastic_count > 0:
                    elastic_share = (elastic_revenue / total_analyzed_revenue * 100) if total_analyzed_revenue > 0 else 0
                    pricing_recommendations.append(
                        f"🔴 **Эластичные товары ({elastic_count} шт., {elastic_share:.1f}% выручки)**: "
                        f"Снижение цены на 10-15% может увеличить объем продаж на >10%. "
                        f"Используйте акции и промо для роста выручки."
                    )
                
                if inelastic_count > 0:
                    inelastic_share = (inelastic_revenue / total_analyzed_revenue * 100) if total_analyzed_revenue > 0 else 0
                    pricing_recommendations.append(
                        f"🟢 **Неэластичные товары ({inelastic_count} шт., {inelastic_share:.1f}% выручки)**: "
                        f"Повышение цены на 5-10% не повлияет критично на спрос. "
                        f"Можно увеличить маржинальность."
                    )
                
                if unit_count > 0:
                    pricing_recommendations.append(
                        f"🟡 **Единично-эластичные товары ({unit_count} шт.)**: "
                        f"Цена близка к оптимальной. Сфокусируйтесь на удержании позиций."
                    )
                
                for rec in pricing_recommendations:
                    st.markdown(f'<div class="insight-card">{rec}</div>', unsafe_allow_html=True)
                
                # Общие выводы по эластичности
                st.info(
                    f"💡 **Ключевой вывод**: Из {len(elasticity_df)} проанализированных товаров "
                    f"{elastic_count} являются эластичными (чувствительны к цене), "
                    f"{inelastic_count} - неэластичными (нечувствительны к цене). "
                    f"Используйте эти данные для оптимизации ценовой стратегии."
                )
                
            else:
                st.warning("⚠️ Недостаточно данных для анализа эластичности спроса. Требуется больше исторических данных с вариацией цен.")
            
            st.markdown("---")
            st.markdown("## 📥 Экспорт результатов")
            
            col1, col2 = st.columns(2)
            
            with col1:
                export_data = detailed_forecast.copy()
                export_data['Магазин'] = selected_magazin
                export_data['Сегмент'] = selected_segment
                
                csv = export_data.to_csv(index=False)
                st.download_button(
                    label="📊 Скачать прогноз (CSV)",
                    data=csv,
                    file_name=f"forecast_{selected_magazin}_{selected_segment}_{forecast_days}days.csv",
                    mime="text/csv",
                    use_container_width=True
                )
            
            with col2:
                # Word отчет
                word_data = create_word_report(
                    detailed_forecast, selected_magazin, selected_segment, forecast_days,
                    total_forecast, avg_daily_forecast, forecast_revenue, confidence_score,
                    accuracy_metrics, insights, filtered_df, prophet_data
                )
                
                if word_data:
                    st.download_button(
                        label="📄 Скачать отчет (WORD)",
                        data=word_data,
                        file_name=f"report_{selected_magazin}_{selected_segment}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                else:
                    st.info("Word недоступен. Установите: pip install python-docx")

if __name__ == "__main__":
    main()
